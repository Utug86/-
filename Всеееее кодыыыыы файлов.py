# код - main.py

import asyncio
import sys
from pathlib import Path
from typing import Optional, Dict, Any, Set
import signal
from datetime import datetime, timedelta
import logging
from dataclasses import dataclass, asdict
import json

from logs import get_logger
from rag_chunk_tracker import ChunkUsageTracker
from rag_retriever import HybridRetriever
from rag_telegram import TelegramPublisher
from rag_lmclient import LMClient
from rag_langchain_tools import enrich_context_with_tools
from rag_prompt_utils import get_prompt_parts
from image_utils import prepare_media_for_post, get_media_type
from utils.config_manager import ConfigManager
from utils.state_manager import StateManager
from utils.exceptions import (
    RAGException,
    ConfigurationError,
    InitializationError,
    ProcessingError,
    ModelError,
    TelegramError,
    FileOperationError,
)

@dataclass
class SystemStats:
    """Статистика работы системы"""
    total_topics: int = 0
    processed_topics: int = 0
    failed_topics: int = 0
    start_time: Optional[datetime] = None
    current_topic: Optional[str] = None
    last_error: Optional[str] = None
    last_processing_time: Optional[float] = None
    total_chars_generated: int = 0
    avg_chars_per_topic: float = 0.0

    def to_dict(self) -> Dict[str, Any]:
        stats = asdict(self)
        if self.start_time:
            stats['start_time'] = self.start_time.isoformat()
            stats['running_time'] = str(datetime.now() - self.start_time)
        stats['success_rate'] = (
            (self.processed_topics / self.total_topics * 100)
            if self.total_topics > 0 else 0
        )
        return stats

class RAGSystem:
    def __init__(self):
        # --- Загрузка конфигурации через ConfigManager ---
        self.config_manager = ConfigManager(Path("config/config.json"))
        self.logger = get_logger(__name__, logfile=self.config_manager.get_path("log_dir") / "bot.log")

        # --- Пути из config_manager ---
        self.data_dir = self.config_manager.get_path("data_dir")
        self.log_dir = self.config_manager.get_path("log_dir")
        self.inform_dir = self.config_manager.get_path("inform_dir")
        self.config_dir = Path("config")  # для совместимости, можно убрать если не требуется
        self.media_dir = self.config_manager.get_path("media_dir")
        self.topics_file = self.data_dir / "topics.txt"
        self.processed_topics_file = self.config_manager.get_path("processed_topics_file")
        self.index_file = self.config_manager.get_path("index_file")
        self.context_file = self.config_manager.get_path("context_file")
        self.usage_stats_file = self.config_manager.get_path("usage_stats_file")

        # --- Статистика ---
        self.stats = SystemStats()

        # --- StateManager для прогресса ---
        self.state_manager = StateManager(self.processed_topics_file)

        # --- Флаг для graceful shutdown ---
        self.should_exit = False
        signal.signal(signal.SIGINT, self.handle_shutdown)
        signal.signal(signal.SIGTERM, self.handle_shutdown)

    def setup_paths(self):
        """Проверка и создание необходимых директорий"""
        required_dirs = [
            self.data_dir,
            self.log_dir,
            self.inform_dir,
            self.config_dir,
            self.media_dir,
            self.data_dir / "prompt_1",
            self.data_dir / "prompt_2"
        ]
        for directory in required_dirs:
            try:
                directory.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                raise InitializationError(f"Failed to create directory {directory}: {e}")

        # Проверка критичных файлов
        if not self.topics_file.exists():
            raise ConfigurationError("topics.txt not found")

    def load_processed_topics(self) -> Set[str]:
        """Загрузка списка обработанных тем через state_manager"""
        return self.state_manager.get_processed_topics()

    def save_processed_topic(self, topic: str):
        """Сохранение обработанной темы через state_manager"""
        self.state_manager.add_processed_topic(topic)

    def add_failed_topic(self, topic: str, error: str):
        """Сохранение темы с ошибкой через state_manager"""
        self.state_manager.add_failed_topic(topic, error)

    async def notify_error(self, message: str):
        """Отправка уведомления об ошибке в Telegram"""
        if hasattr(self, 'telegram'):
            try:
                await self.telegram.send_text(
                    f"🚨 RAG System Error:\n{message}"
                )
            except Exception as e:
                self.logger.error(f"Failed to send error notification: {e}")

    def handle_shutdown(self, signum, frame):
        """Обработчик сигналов завершения"""
        self.logger.info("Received shutdown signal, cleaning up...")
        self.should_exit = True

        # Сохранение статистики перед выходом
        try:
            stats = self.stats.to_dict()
            stats_file = self.log_dir / f"stats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(stats_file, 'w', encoding='utf-8') as f:
                json.dump(stats, f, indent=2, ensure_ascii=False)
            self.logger.info(f"Statistics saved to {stats_file}")
        except Exception as e:
            self.logger.error(f"Failed to save statistics: {e}")

    def _load_remaining_topics(self) -> list:
        """Загрузка оставшихся тем для обработки"""
        try:
            all_topics = self.topics_file.read_text(encoding='utf-8').splitlines()
            processed = self.load_processed_topics()
            remaining = [t for t in all_topics if t not in processed]
            self.logger.info(f"Loaded {len(remaining)} remaining topics")
            return remaining
        except Exception as e:
            raise ProcessingError(f"Failed to load topics: {e}")

        async def process_topics(self):
        """Обработка тем из topics.txt"""
        topics = self._load_remaining_topics()
        self.stats.total_topics = len(topics)
        self.stats.start_time = datetime.now()

        for topic in topics:
            if self.should_exit:
                break
            self.stats.current_topic = topic
            processing_start = datetime.now()
            try:
                self.logger.info(
                    f"Processing topic {self.stats.processed_topics + 1}/{self.stats.total_topics}: {topic}"
                )
                text_length = await self.process_single_topic(topic)
                self.stats.processed_topics += 1
                self.stats.total_chars_generated += text_length
                self.stats.avg_chars_per_topic = (
                    self.stats.total_chars_generated / self.stats.processed_topics
                )
                self.save_processed_topic(topic)
                self.stats.last_processing_time = (
                    datetime.now() - processing_start
                ).total_seconds()
            except ProcessingError as e:
                error_msg = f"ProcessingError: {e}"
                self.logger.error(error_msg, exc_info=True)
                self.stats.failed_topics += 1
                self.stats.last_error = error_msg
                self.add_failed_topic(topic, error_msg)
                await self.notify_error(error_msg)
                await asyncio.sleep(5)
            except Exception as e:
                error_msg = f"Unexpected error processing topic {topic}: {e}"
                self.logger.critical(error_msg, exc_info=True)
                self.stats.failed_topics += 1
                self.stats.last_error = error_msg
                self.add_failed_topic(topic, error_msg)
                await self.notify_error(error_msg)
                await asyncio.sleep(5)

    async def process_single_topic(self, topic: str) -> int:
        """
        Обработка одной темы.
        Возвращает длину сгенерированного текста.
        """
        try:
            # Получение контекста из RAG
            context = self.retriever.retrieve(topic)

            # Обогащение контекста дополнительными инструментами
            context = enrich_context_with_tools(topic, context, self.inform_dir)

            # Поиск файлов промптов
            prompt1_files = sorted((self.data_dir / "prompt_1").glob("*.txt"))
            prompt2_files = sorted((self.data_dir / "prompt_2").glob("*.txt"))

            if not prompt1_files or not prompt2_files:
                raise ProcessingError("No prompt files found")

            import random
            file1 = random.choice(prompt1_files)
            file2 = random.choice(prompt2_files)

            # Генерация промпта
            prompt_full = get_prompt_parts(
                data_dir=self.data_dir,
                topic=topic,
                context=context,
                file1=file1,
                file2=file2
            )

            max_chars = (
                self.llm_config["max_chars_with_media"]
                if "{UPLOADFILE}" in prompt_full
                else self.llm_config["max_chars"]
            )

            # Генерация текста
            text = await self.lm.generate(
                topic,
                max_chars=max_chars
            )

            if not text:
                raise ProcessingError("Failed to generate text")

            # Отправка в Telegram
            if "{UPLOADFILE}" in prompt_full:
                await self.handle_media_post(text)
            else:
                await self.telegram.send_text(text)

            self.logger.info(
                f"Successfully processed topic: {topic}, "
                f"text length: {len(text)}"
            )

            return len(text)

        except Exception as e:
            raise ProcessingError(f"Failed to process topic {topic}: {e}")

    async def handle_media_post(self, text: str):
        """Обработка поста с медиафайлом"""
        try:
            media_file = prepare_media_for_post(self.media_dir)
            if not media_file:
                raise ProcessingError("No valid media file found")

            media_type = get_media_type(media_file)
            self.logger.info(f"Selected media file: {media_file} (type: {media_type})")

            media_handlers = {
                "image": self.telegram.send_photo,
                "video": self.telegram.send_video,
                "document": self.telegram.send_document,
                "audio": self.telegram.send_audio
            }

            if media_type in media_handlers:
                await media_handlers[media_type](media_file, caption=text)
            else:
                self.logger.warning(f"Unknown media type: {media_file}")
                await self.telegram.send_text(text)

        except Exception as e:
            self.logger.error(f"Media handling error: {e}")
            await self.telegram.send_text(text)

    async def run(self):
        """Основной метод запуска системы"""
        try:
            # --- Загрузка параметров из конфиг-менеджера ---
            self.telegram_config = self.config_manager.get_telegram_config()
            self.llm_config = self.config_manager.get_llm_config()
            self.retrieval_config = self.config_manager.get_retrieval_config()
            self.system_config = self.config_manager.get_system_config()

            # --- Инициализация компонентов ---
            self.setup_paths()

            self.usage_tracker = ChunkUsageTracker(
                usage_stats_file=self.usage_stats_file,
                logger=self.logger,
                chunk_usage_limit=self.system_config["chunk_usage_limit"],
                usage_reset_days=self.system_config["usage_reset_days"],
                diversity_boost=self.system_config["diversity_boost"]
            )
            self.usage_tracker.cleanup_old_stats()

            self.retriever = HybridRetriever(
                emb_model=self.retrieval_config["embedding_model"],
                cross_model=self.retrieval_config["cross_encoder"],
                index_file=self.index_file,
                context_file=self.context_file,
                inform_dir=self.inform_dir,
                chunk_size=self.retrieval_config["chunk_size"],
                overlap=self.retrieval_config["overlap"],
                top_k_title=self.retrieval_config["top_k_title"],
                top_k_faiss=self.retrieval_config["top_k_faiss"],
                top_k_final=self.retrieval_config["top_k_final"],
                usage_tracker=self.usage_tracker,
                logger=self.logger
            )

            self.lm = LMClient(
                retriever=self.retriever,
                data_dir=self.data_dir,
                inform_dir=self.inform_dir,
                logger=self.logger,
                model_url=self.llm_config["url"],
                model_name=self.llm_config["model_name"],
                max_tokens=self.llm_config["max_tokens"],
                max_chars=self.llm_config["max_chars"],
                temperature=self.llm_config["temperature"],
                timeout=self.llm_config["timeout"],
                history_lim=self.llm_config["history_limit"],
                system_msg=self.llm_config.get("system_message")
            )

            self.telegram = TelegramPublisher(
                self.telegram_config["bot_token"],
                self.telegram_config["channel_id"],
                logger=self.logger,
                max_retries=self.telegram_config["retry_attempts"],
                retry_delay=self.telegram_config["retry_delay"],
                enable_preview=self.telegram_config["enable_preview"]
            )

            # Проверка соединения с Telegram
            if not await self.telegram.check_connection():
                raise TelegramError("Failed to connect to Telegram")

            self.logger.info("System initialized successfully")

            # Основной цикл обработки
            await self.process_topics()

        except ConfigurationError as e:
            self.logger.critical(f"Configuration error: {e}")
            await self.notify_error(f"Configuration error: {e}")
            sys.exit(1)
        except Exception as e:
            self.logger.critical(f"Unexpected error: {e}")
            await self.notify_error(f"Unexpected error: {e}")
            sys.exit(1)
        finally:
            # Сохранение статистики и очистка
            if hasattr(self, 'usage_tracker'):
                self.usage_tracker.save_statistics()

            # Вывод итоговой статистики
            stats = self.stats.to_dict()
            self.logger.info("Final statistics:")
            for key, value in stats.items():
                self.logger.info(f"{key}: {value}")

            self.logger.info("System shutdown complete")

def main():
    """Точка входа с обработкой всех возможных ошибок"""
    def handle_exception(exc_type, exc_value, exc_traceback):
        if issubclass(exc_type, KeyboardInterrupt):
            sys.__excepthook__(exc_type, exc_value, exc_traceback)
            return
        logging.error(
            "Uncaught exception",
            exc_info=(exc_type, exc_value, exc_traceback)
        )
    sys.excepthook = handle_exception

    rag_system = RAGSystem()
    try:
        asyncio.run(rag_system.run())
    except KeyboardInterrupt:
        print("\nShutdown requested... exiting")
    except Exception as e:
        print(f"Fatal error: {e}")
        sys.exit(1)

if __name__ == '__main__':
    main()

# код - logs.py

import logging

def get_logger(name: str, logfile: str = None, level=logging.INFO) -> logging.Logger:
    logger = logging.getLogger(name)
    if not logger.hasHandlers():
        if logfile:
            fh = logging.FileHandler(logfile, encoding="utf-8")
            fh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
            logger.addHandler(fh)
        sh = logging.StreamHandler()
        sh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
        logger.addHandler(sh)
        logger.setLevel(level)
    return logger

# код - rag_file_utils.py

from pathlib import Path
import logging
from typing import Optional
import pandas as pd
from bs4 import BeautifulSoup

from logs import get_logger

logger = get_logger("rag_file_utils")

# --- Dynamic imports ---
def _try_import_docx() -> Optional[object]:
    try:
        import docx
        return docx
    except ImportError:
        logger.warning("python-docx не установлен. Форматы .docx будут проигнорированы.")
        return None

def _try_import_pypdf2() -> Optional[object]:
    try:
        import PyPDF2
        return PyPDF2
    except ImportError:
        logger.warning("PyPDF2 не установлен. Форматы .pdf будут проигнорированы.")
        return None

def _try_import_textract() -> Optional[object]:
    try:
        import textract
        return textract
    except ImportError:
        logger.warning("textract не установлен. Некоторые форматы могут быть не поддержаны.")
        return None

DOCX = _try_import_docx()
PDF = _try_import_pypdf2()
TEXTRACT = _try_import_textract()

COMMON_ENCODINGS = ["utf-8", "cp1251", "windows-1251", "latin-1", "iso-8859-1"]
MAX_FILE_SIZE_MB = 100

def _smart_read_text(path: Path) -> str:
    """
    Пробует прочитать текстовый файл с помощью популярных кодировок.
    Возвращает содержимое файла или пустую строку при ошибке.
    """
    for encoding in COMMON_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except Exception as e:
            logger.debug(f"Проблема с кодировкой {encoding} для {path}: {e}")
    logger.error(f"Не удалось прочитать файл {path} в поддерживаемых кодировках: {COMMON_ENCODINGS}")
    return ""

def extract_text_from_file(path: Path) -> str:
    """
    Универсальный парсер для извлечения текста из файлов различных форматов.
    Поддерживает: txt, html, csv, xlsx, xlsm, docx, doc, pdf.
    Возвращает текст или специальную метку при ошибке/неподдерживаемом формате.
    """
    ext = path.suffix.lower()
    if not path.exists():
        logger.error(f"Файл не найден: {path}")
        return "[Файл не найден]"

    if path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        logger.warning(f"Файл слишком большой (> {MAX_FILE_SIZE_MB} МБ): {path}")
        return "[Файл слишком большой для обработки]"

    try:
        if ext == ".txt":
            logger.info(f"Extracting text from TXT file: {path}")
            return _smart_read_text(path)

        elif ext == ".html":
            logger.info(f"Extracting text from HTML file: {path}")
            html_content = _smart_read_text(path)
            soup = BeautifulSoup(html_content, "html.parser")
            return soup.get_text(separator=" ")

        elif ext == ".csv":
            logger.info(f"Extracting text from CSV file: {path}")
            try:
                df = pd.read_csv(path)
                return df.to_csv(sep="\t", index=False)
            except Exception as e:
                logger.warning(f"Ошибка чтения CSV через pandas: {e}. Пробуем как текст.")
                return _smart_read_text(path)

        elif ext in [".xlsx", ".xls", ".xlsm"]:
            logger.info(f"Extracting text from Excel file: {path}")
            try:
                df = pd.read_excel(path)
                return df.to_csv(sep="\t", index=False)
            except Exception as e:
                logger.error(f"Ошибка чтения Excel через pandas: {e}")
                return "[Ошибка чтения Excel]"

        elif ext == ".docx":
            logger.info(f"Extracting text from DOCX file: {path}")
            if DOCX is not None:
                try:
                    doc = DOCX.Document(path)
                    return "\n".join([p.text for p in doc.paragraphs])
                except Exception as e:
                    logger.error(f"Ошибка чтения DOCX: {e}")
                    return "[Ошибка чтения DOCX]"
            else:
                logger.warning(f"Модуль python-docx не установлен. DOCX не поддерживается.")
                return "[Формат DOCX не поддерживается]"

        elif ext == ".doc":
            logger.info(f"Extracting text from DOC file: {path}")
            if TEXTRACT is not None:
                try:
                    return TEXTRACT.process(str(path)).decode("utf-8", errors="ignore")
                except Exception as e:
                    logger.error(f"Ошибка чтения DOC через textract: {e}")
                    return "[Ошибка чтения DOC]"
            else:
                logger.warning(f"textract не установлен. DOC не поддерживается.")
                return "[Формат DOC не поддерживается]"

        elif ext == ".pdf":
            logger.info(f"Extracting text from PDF file: {path}")
            if PDF is not None:
                try:
                    text = []
                    with open(path, "rb") as f:
                        reader = PDF.PdfReader(f)
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text.append(page_text)
                    return "\n".join(text)
                except Exception as e:
                    logger.warning(f"Ошибка чтения PDF через PyPDF2: {e}. Пробуем textract.")
            if TEXTRACT is not None:
                try:
                    return TEXTRACT.process(str(path)).decode("utf-8", errors="ignore")
                except Exception as e:
                    logger.error(f"Ошибка чтения PDF через textract: {e}")
                    return "[Ошибка чтения PDF]"
            logger.warning(f"PyPDF2 и textract не установлены. PDF не поддерживается.")
            return "[Формат PDF не поддерживается]"

        else:
            logger.warning(f"Неподдерживаемый тип файла: {path}")
            return f"[Неподдерживаемый тип файла: {ext}]"

    except Exception as e:
        logger.error(f"Критическая ошибка при извлечении текста из {path}: {e}")
        return "[Критическая ошибка при извлечении текста]"

def clean_html_from_cell(cell_value) -> str:
    """
    Очищает строку/ячейку от HTML-тегов.
    """
    if isinstance(cell_value, str):
        return BeautifulSoup(cell_value, "html.parser").get_text(separator=" ")
    return str(cell_value)

# код - rag_chunk_tracker.py

import json
import hashlib
from collections import Counter, deque
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Callable, Optional, Tuple

class ChunkUsageTracker:
    """
    Трекер использования чанков — хранит, обновляет и очищает статистику использования фрагментов знаний/контекста.
    Поддерживает устойчивые идентификаторы чанков (hash чанка), версионность базы, гибкие penalty/boost функции,
    полную очистку по возрасту и расширенные методы аналитики.
    """

    def __init__(
        self,
        usage_stats_file: Path,
        logger,
        chunk_usage_limit: int,
        usage_reset_days: int,
        diversity_boost: float,
        index_version: Optional[str] = None,
        index_hash: Optional[str] = None,
        penalty_func: Optional[Callable[[int, int, int], float]] = None,
        boost_func: Optional[Callable[[int, int], float]] = None,
    ):
        """
        :param usage_stats_file: Путь к файлу статистики использования
        :param logger: Логгер
        :param chunk_usage_limit: Лимит использования для penalty
        :param usage_reset_days: Сколько дней хранить usage перед очисткой
        :param diversity_boost: Базовый коэффициент diversity
        :param index_version: Версия индекса/базы знаний (для сброса статистики при обновлении)
        :param index_hash: Хеш базы знаний (для сброса статистики при обновлении)
        :param penalty_func: Кастомная функция penalty, принимает (chunk_count, title_count, chunk_usage_limit)
        :param boost_func: Кастомная функция diversity boost, принимает (chunk_count, chunk_usage_limit)
        """
        self.usage_stats_file: Path = usage_stats_file
        self.logger = logger
        self.chunk_usage_limit = chunk_usage_limit
        self.usage_reset_days = usage_reset_days
        self.diversity_boost = diversity_boost
        self.index_version = index_version
        self.index_hash = index_hash

        self.penalty_func = penalty_func or self._default_penalty_func
        self.boost_func = boost_func or self._default_boost_func

        self.usage_stats: Dict[str, Dict[str, Any]] = {}
        self.recent_usage: deque = deque(maxlen=100)
        self.title_usage: Counter = Counter()
        self.chunk_usage: Counter = Counter()
        self.session_usage: Counter = Counter()
        self.loaded_index_version = None
        self.loaded_index_hash = None
        self.load_statistics()

    @staticmethod
    def get_chunk_hash(chunk_text: str, source: Optional[str]=None) -> str:
        """
        Возвращает устойчивый идентификатор чанка (sha1 от текста + source).
        """
        to_hash = (chunk_text or "") + "|" + (source or "")
        return hashlib.sha1(to_hash.encode('utf-8')).hexdigest()

    def load_statistics(self):
        """
        Загружает статистику из файла. Если версия/хеш базы отличается — сбрасывает usage.
        """
        try:
            if self.usage_stats_file.exists():
                with open(self.usage_stats_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.usage_stats = data.get('usage_stats', {})
                    self.title_usage = Counter(data.get('title_usage', {}))
                    self.chunk_usage = Counter(data.get('chunk_usage', {}))
                    self.recent_usage = deque(data.get('recent_usage', []), maxlen=100)
                    self.loaded_index_version = data.get('index_version')
                    self.loaded_index_hash = data.get('index_hash')
                if (self.index_version and self.loaded_index_version != self.index_version) or \
                   (self.index_hash and self.loaded_index_hash != self.index_hash):
                    self.logger.warning("Knowledge base index version/hash mismatch: usage statistics will be reset.")
                    self.clear_all_statistics()
                else:
                    self.logger.info("Loaded usage statistics successfully")
        except Exception as e:
            self.logger.warning(f"Failed to load usage statistics: {e}")
            self.usage_stats = {}

    def save_statistics(self):
        """
        Сохраняет статистику в файл (atomic save).
        """
        try:
            data = {
                'usage_stats': self.usage_stats,
                'title_usage': dict(self.title_usage),
                'chunk_usage': dict(self.chunk_usage),
                'recent_usage': list(self.recent_usage),
                'last_updated': datetime.now().isoformat(),
                'index_version': self.index_version,
                'index_hash': self.index_hash
            }
            tmp_file = self.usage_stats_file.with_suffix('.tmp')
            with open(tmp_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            tmp_file.replace(self.usage_stats_file)
        except Exception as e:
            self.logger.error(f"Failed to save usage statistics: {e}")

    def record_usage(self, chunk_hashes: List[str], titles: List[str], metadata: List[Dict[str, Any]]):
        """
        Записывает использование чанков (по hash'ам).
        :param chunk_hashes: Список хешей чанков (уникальные id)
        :param titles: Список тайтлов источников (соответствует chunk_hashes)
        :param metadata: Массив метадаты (может быть пустой, используется для расширения)
        """
        timestamp = datetime.now().isoformat()
        for i, chunk_hash in enumerate(chunk_hashes):
            title = titles[i] if i < len(titles) else "unknown"
            if chunk_hash not in self.usage_stats:
                self.usage_stats[chunk_hash] = {
                    'count': 0,
                    'last_used': None,
                    'title': title
                }
            self.usage_stats[chunk_hash]['count'] += 1
            self.usage_stats[chunk_hash]['last_used'] = timestamp
            self.title_usage[title] += 1
            self.chunk_usage[chunk_hash] += 1
            self.session_usage[chunk_hash] += 1
            self.recent_usage.append({
                'chunk_id': chunk_hash,
                'title': title,
                'timestamp': timestamp
            })
        self.save_statistics()

    def get_usage_penalty(self, chunk_hash: str, title: str) -> float:
        """
        Возвращает штраф за частое использование (0.0 - 1.5), с учетом настраиваемой penalty-функции.
        """
        chunk_count = self.usage_stats.get(chunk_hash, {}).get('count', 0)
        title_count = self.title_usage.get(title, 0)
        return self.penalty_func(chunk_count, title_count, self.chunk_usage_limit)

    def get_diversity_boost(self, chunk_hash: str, title: str) -> float:
        """
        Возвращает буст для редко используемых чанков, с учетом настраиваемой boost-функции.
        """
        chunk_count = self.usage_stats.get(chunk_hash, {}).get('count', 0)
        return self.boost_func(chunk_count, self.chunk_usage_limit)

    @staticmethod
    def _default_penalty_func(chunk_count: int, title_count: int, chunk_usage_limit: int) -> float:
        """
        Стандартная penalty-функция (можно заменить через __init__).
        """
        chunk_penalty = min(chunk_count / chunk_usage_limit, 1.0)
        title_penalty = min(title_count / (chunk_usage_limit * 2), 0.5)
        return chunk_penalty + title_penalty

    @staticmethod
    def _default_boost_func(chunk_count: int, chunk_usage_limit: int) -> float:
        """
        Стандартная функция diversity boost (можно заменить через __init__).
        """
        if chunk_count == 0:
            return 2.0
        elif chunk_count < chunk_usage_limit // 3:
            return 1.0
        else:
            return 0.0

    def cleanup_old_stats(self, full_reset: bool = False):
        """
        Очищает старую статистику по времени, либо полностью сбрасывает usage по возрасту.
        :param full_reset: Если True — полностью сбрасывает usage у всех чанков, last_used которых старше порога.
        """
        cutoff_date = datetime.now() - timedelta(days=self.usage_reset_days)
        cutoff_str = cutoff_date.isoformat()
        cleaned_count = 0
        for chunk_hash in list(self.usage_stats.keys()):
            last_used = self.usage_stats[chunk_hash].get('last_used')
            if last_used and last_used < cutoff_str:
                if full_reset:
                    del self.usage_stats[chunk_hash]
                    cleaned_count += 1
                else:
                    old_count = self.usage_stats[chunk_hash]['count']
                    self.usage_stats[chunk_hash]['count'] = max(0, old_count - 1)
                    if self.usage_stats[chunk_hash]['count'] == 0:
                        del self.usage_stats[chunk_hash]
                        cleaned_count += 1
        if cleaned_count > 0:
            self.logger.info(f"Cleaned up {cleaned_count} old usage statistics entries{' (full reset)' if full_reset else ''}")
            self.save_statistics()

    def clear_all_statistics(self):
        """
        Полный сброс всей статистики использования.
        """
        self.usage_stats.clear()
        self.title_usage.clear()
        self.chunk_usage.clear()
        self.session_usage.clear()
        self.recent_usage.clear()
        self.save_statistics()
        self.logger.info("All usage statistics cleared.")

    def get_unused_chunks(self, metadata: List[Dict[str, Any]]) -> List[str]:
        """
        Возвращает список hash'ей неиспользованных чанков из metadata.
        """
        all_hashes = set(self.get_chunk_hash(m['chunk'], m.get('source')) for m in metadata)
        used = set(self.usage_stats.keys())
        return list(all_hashes - used)

    def get_top_used_chunks(self, n: int = 10) -> List[Tuple[str, int]]:
        """
        Возвращает топ-N наиболее часто используемых чанков (hash, count).
        """
        return self.chunk_usage.most_common(n)

    def get_usage_distribution(self) -> Dict[str, int]:
        """
        Возвращает распределение использования по чанкам (hash -> count).
        """
        return dict(self.chunk_usage)

    def get_stats_summary(self) -> Dict[str, Any]:
        """
        Краткая сводка по статистике использования.
        """
        total_chunks = len(self.usage_stats)
        total_titles = len(self.title_usage)
        never_used = sum(1 for data in self.usage_stats.values() if data.get('count', 0) == 0)
        most_used = self.chunk_usage.most_common(1)[0] if self.chunk_usage else ("", 0)
        return {
            "total_chunks": total_chunks,
            "total_titles": total_titles,
            "never_used_chunks": never_used,
            "most_used_chunk": most_used,
        }

    # Документация property-методов
    @property
    def usage_stats_count(self) -> int:
        """Текущее число уникальных чанков с usage-статистикой."""
        return len(self.usage_stats)

    @property
    def title_count(self) -> int:
        """Число различных источников (title) в usage-статистике."""
        return len(self.title_usage)

# код - rag_retriever.py

import faiss
import numpy as np
import json
import shutil
from pathlib import Path
from sentence_transformers import SentenceTransformer, CrossEncoder
from typing import Any, Dict, List, Optional, Union
import hashlib
import logging
import datetime

from rag_file_utils import extract_text_from_file

def notify_admin(message: str) -> None:
    logging.warning(f"[ADMIN NOTIFY] {message}")

class HybridRetriever:
    """
    Гибридный ретривер: поиск релевантных чанков по FAISS-индексу и кросс-энкодеру.
    Гарантирует валидацию индекса, atomic запись, резервное копирование, обработку edge-cases.
    """
    INDEX_VERSION = "1.2"

    def __init__(
        self,
        emb_model: str,
        cross_model: str,
        index_file: Path,
        context_file: Path,
        inform_dir: Path,
        chunk_size: int,
        overlap: int,
        top_k_title: int,
        top_k_faiss: int,
        top_k_final: int,
        usage_tracker,
        logger: logging.Logger
    ):
        self.emb_model = emb_model
        self.cross_model = cross_model
        self.index_file = Path(index_file)
        self.context_file = Path(context_file)
        self.inform_dir = Path(inform_dir)
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.top_k_title = top_k_title
        self.top_k_faiss = top_k_faiss
        self.top_k_final = top_k_final
        self.usage_tracker = usage_tracker
        self.logger = logger

        self.sentencemodel = SentenceTransformer(self.emb_model)
        self.crossencoder = CrossEncoder(self.cross_model)
        self.faiss_index: Optional[faiss.Index] = None
        self.metadata: Optional[List[Dict[str, Any]]] = None
        self.index_metadata: Dict[str, Any] = {}

        self._try_load_or_build_indices()

    def _get_index_signature(self) -> str:
        conf = {
            "version": self.INDEX_VERSION,
            "emb_model": self.emb_model,
            "cross_model": self.cross_model,
            "chunk_size": self.chunk_size,
            "overlap": self.overlap,
            "inform_dir_hash": self._dir_hash(self.inform_dir)
        }
        return hashlib.sha256(json.dumps(conf, sort_keys=True).encode()).hexdigest()

    @staticmethod
    def _dir_hash(directory: Path) -> str:
        if not directory.exists():
            return ""
        files = sorted([(str(f), f.stat().st_mtime) for f in directory.iterdir() if f.is_file()])
        return hashlib.sha256(json.dumps(files, sort_keys=True).encode()).hexdigest()

    def _try_load_or_build_indices(self) -> None:
        self.logger.info("Initializing HybridRetriever...")
        rebuild_needed = False
        if self.index_file.exists() and self.context_file.exists():
            try:
                self._load_indices()
                idx_sig = self.index_metadata.get("index_signature")
                expected_sig = self._get_index_signature()
                if idx_sig != expected_sig:
                    self.logger.warning("Index signature mismatch. Rebuilding index...")
                    notify_admin("HybridRetriever: Index signature mismatch, rebuild triggered.")
                    rebuild_needed = True
            except Exception as e:
                self.logger.warning(f"Failed to load indices: {e}. Rebuilding...")
                notify_admin(f"HybridRetriever: Failed to load indices: {e}. Rebuilding.")
                rebuild_needed = True
        else:
            self.logger.info("No existing indices found. Building new ones...")
            rebuild_needed = True
        if rebuild_needed:
            self._build_indices()

    def _load_indices(self) -> None:
        self.logger.info("Loading indices...")
        try:
            with open(self.context_file, 'r', encoding='utf-8') as f:
                context_json = json.load(f)
            if not isinstance(context_json, dict) or "metadata" not in context_json:
                raise ValueError("context_file missing 'metadata' field")
            self.metadata = context_json["metadata"]
            self.index_metadata = context_json.get("index_metadata", {})
            if not isinstance(self.metadata, list) or len(self.metadata) == 0:
                raise ValueError("Metadata must be a non-empty list")
            # Ensure fields exist
            for item in self.metadata:
                item.setdefault('tokens', item['chunk'].split())
                item.setdefault('created_at', None)
                item.setdefault('source', None)
            self.faiss_index = faiss.read_index(str(self.index_file))
            # Validate dimension
            sample_chunk = self.metadata[0]['chunk']
            emb_sample = self.sentencemodel.encode([sample_chunk], convert_to_tensor=False, normalize_embeddings=True)
            if self.faiss_index.d != emb_sample.shape[1]:
                self.logger.critical(
                    f"Embedding dimension mismatch: FAISS index {self.faiss_index.d}, model {emb_sample.shape[1]}"
                )
                notify_admin("HybridRetriever: Embedding dimension mismatch, index rebuild required")
                raise RuntimeError("Embedding dimension mismatch — rebuild index!")
            self.logger.info(f"Loaded {len(self.metadata)} chunks")
            self.logger.info(f"Index metadata: {self.index_metadata}")
        except Exception as e:
            self.logger.critical(f"Error loading indices: {e}")
            notify_admin(f"HybridRetriever: Error loading indices: {e}")
            raise

    def _normalize_text(self, text: str) -> str:
        """
        Базовая нормализация текста (html → plain, lower, clean).
        """
        import re
        from bs4 import BeautifulSoup
        text = BeautifulSoup(text, "html.parser").get_text(separator=" ")
        text = text.lower()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^а-яa-z0-9\s\.,:;!\?\(\)\[\]\'\"-]', '', text)
        return text.strip()

    def _semantic_deduplicate(self, chunks: List[Dict[str, Any]], threshold: float = 0.91) -> List[Dict[str, Any]]:
        """
        Семантическая дедупликация: удаление дубликатов по cosine similarity эмбеддингов.
        Для больших коллекций (N>10000) пропускается.
        """
        if len(chunks) < 2:
            return chunks
        if len(chunks) > 10000:
            self.logger.warning("Skipping semantic deduplication due to excessive collection size")
            return chunks
        texts = [c['chunk'] for c in chunks]
        embs = self.sentencemodel.encode(texts, convert_to_tensor=False, show_progress_bar=False, normalize_embeddings=True)
        embs = np.asarray(embs, dtype='float32')
        to_keep = []
        already_used = set()
        for i in range(len(chunks)):
            if i in already_used:
                continue
            to_keep.append(chunks[i])
            for j in range(i + 1, len(chunks)):
                if j in already_used:
                    continue
                sim = np.dot(embs[i], embs[j])
                if sim > threshold:
                    already_used.add(j)
        self.logger.info(f"Semantic deduplication: {len(chunks)} → {len(to_keep)} unique chunks (threshold={threshold})")
        return to_keep

    def _build_indices(self) -> None:
        self.logger.info("Building new indices...")
        metadata = []
        inform_files = [f for f in self.inform_dir.iterdir()
                        if f.suffix.lower() in [".txt", ".html", ".csv", ".xlsx", ".xlsm", ".doc", ".docx", ".pdf"]]
        if not inform_files:
            notify_admin(f"HybridRetriever: No suitable files in {self.inform_dir}")
            raise RuntimeError(f"No suitable files in {self.inform_dir}")
        self.logger.info(f"Found {len(inform_files)} files to process")
        index_time = datetime.datetime.utcnow().isoformat()
        for file in inform_files:
            try:
                title = file.stem.lower()
                text = extract_text_from_file(file)
                if not text or not text.strip():
                    self.logger.warning(f"Empty or unreadable file: {file}")
                    continue
                text = self._normalize_text(text)
                words = text.split()
                chunks = []
                for i in range(0, len(words), self.chunk_size - self.overlap):
                    chunk = ' '.join(words[i:i + self.chunk_size])
                    if len(chunk.strip()) < 10:
                        continue
                    tokens = chunk.split()
                    if any(chunk == m['chunk'] for m in chunks):
                        continue
                    chunks.append({'title': title, 'chunk': chunk, 'tokens': tokens,
                                   'created_at': index_time, 'source': str(file)})
                if not chunks:
                    continue
                deduped_chunks = self._semantic_deduplicate(chunks, threshold=0.91)
                metadata.extend(deduped_chunks)
                self.logger.info(f"Processed {file.name}: {len(words)} words -> {len(deduped_chunks)} unique chunks")
            except Exception as e:
                self.logger.error(f"Error processing file {file}: {e}")
                notify_admin(f"HybridRetriever: Error processing file {file}: {e}")
        if not metadata:
            notify_admin("HybridRetriever: No valid chunks created from files")
            raise RuntimeError("No valid chunks created from files")
        metadata = self._semantic_deduplicate(metadata, threshold=0.91)
        self.logger.info(f"Total unique chunks after global deduplication: {len(metadata)}")
        try:
            # Backup before overwrite
            if self.index_file.exists():
                shutil.copy2(self.index_file, self.index_file.with_suffix('.bak'))
            if self.context_file.exists():
                shutil.copy2(self.context_file, self.context_file.with_suffix('.bak'))
            texts = [f"{m['title']}: {m['chunk']}" for m in metadata]
            self.logger.info("Generating embeddings...")
            embs = self.sentencemodel.encode(texts, convert_to_tensor=False, show_progress_bar=True, batch_size=32, normalize_embeddings=True)
            embs = np.asarray(embs, dtype='float32')
            dim = embs.shape[1]
            if len(metadata) > 10000:
                self.logger.info("Large dataset: using HNSW index.")
                faiss_index = faiss.IndexHNSWFlat(dim, 32)
            else:
                faiss_index = faiss.IndexFlatL2(dim)
            faiss_index.add(embs)
            # Atomic write
            tmp_index = self.index_file.with_suffix('.tmp')
            faiss.write_index(faiss_index, str(tmp_index))
            tmp_index.replace(self.index_file)
            self.faiss_index = faiss_index
            self.index_metadata = {
                "index_signature": self._get_index_signature(),
                "version": self.INDEX_VERSION,
                "emb_model": self.emb_model,
                "cross_model": self.cross_model,
                "chunk_size": self.chunk_size,
                "overlap": self.overlap,
                "created_at": index_time,
                "num_chunks": len(metadata),
                "inform_dir": str(self.inform_dir),
            }
            context_json = {
                "metadata": metadata,
                "index_metadata": self.index_metadata
            }
            tmp_context = self.context_file.with_suffix('.tmp')
            with open(tmp_context, 'w', encoding='utf-8') as f:
                json.dump(context_json, f, ensure_ascii=False, indent=2)
            tmp_context.replace(self.context_file)
            self.metadata = metadata
            self.logger.info(f"Indices built and saved: {len(metadata)} unique chunks, index type: {type(self.faiss_index)}")
            self.logger.info(f"Index metadata: {self.index_metadata}")
            notify_admin(f"HybridRetriever: Index rebuilt. {len(metadata)} unique chunks.")
        except Exception as e:
            self.logger.error(f"Error building or saving indices: {e}")
            notify_admin(f"HybridRetriever: Error building index: {e}")
            raise

    def retrieve(self, query: str, return_chunks: bool = False) -> Union[str, List[Dict[str, Any]]]:
        """
        Возвращает релевантный контекст (или чанки) по запросу.
        :param query: Запрос пользователя
        :param return_chunks: Если True, возвращает список чанков, иначе — строку
        """
        self.logger.info(f"Retrieving context for query: '{query}'")
        if self.faiss_index is None or self.metadata is None or len(self.metadata) == 0:
            self.logger.error("Index not loaded or metadata is empty")
            notify_admin("HybridRetriever: Retrieval failed — index not loaded or metadata is empty")
            return "" if not return_chunks else []
        query_emb = self.sentencemodel.encode([query], convert_to_tensor=False, normalize_embeddings=True)
        query_emb = np.asarray(query_emb, dtype='float32')
        if self.faiss_index.d != query_emb.shape[1]:
            self.logger.critical(
                f"Embedding dimension mismatch: FAISS index {self.faiss_index.d}, query_emb {query_emb.shape[1]}"
            )
            notify_admin("HybridRetriever: Embedding dimension mismatch, index rebuild required")
            raise RuntimeError("Embedding dimension mismatch — rebuild index!")
        D, I = self.faiss_index.search(query_emb, self.top_k_faiss)
        I = I[0]
        D = D[0]
        if not len(I):
            self.logger.warning("No relevant chunks found in index")
            return "" if not return_chunks else []
        candidates = []
        for idx, dist in zip(I, D):
            if idx < 0 or idx >= len(self.metadata):
                continue
            meta = self.metadata[idx].copy()
            meta['faiss_dist'] = float(dist)
            candidates.append(meta)
        for c in candidates:
            c['usage_penalty'] = self.usage_tracker.get_usage_penalty(
                c['chunk'], c['title']) if self.usage_tracker else 0.0
        candidates.sort(key=lambda x: (x['faiss_dist'] + x['usage_penalty']))
        rerank_candidates = candidates[:self.top_k_final * 2]
        ce_scores = []
        try:
            if self.crossencoder:
                pairs = [[query, c['chunk']] for c in rerank_candidates]
                ce_scores = self.crossencoder.predict(pairs)
                for c, score in zip(rerank_candidates, ce_scores):
                    c['cross_score'] = float(score)
                rerank_candidates.sort(key=lambda x: -x.get('cross_score', 0))
            else:
                for c in rerank_candidates:
                    c['cross_score'] = 0.0
        except Exception as e:
            self.logger.error(f"Cross-encoder rerank failed: {e}")
            for c in rerank_candidates:
                c['cross_score'] = 0.0
        selected = []
        titles = set()
        for c in rerank_candidates:
            if len(selected) >= self.top_k_final:
                break
            if c['title'] not in titles or self.top_k_final > len(rerank_candidates):
                selected.append(c)
                titles.add(c['title'])
        result = "\n\n".join([f"[{c['title']}] {c['chunk']}" for c in selected])
        self.logger.info(f"Retrieved {len(selected)} chunks from index for query '{query}'")
        return selected if return_chunks else result

    def get_index_stats(self) -> Dict[str, Any]:
        """
        Возвращает статистику по индексу.
        """
        stats = {
            "num_chunks": len(self.metadata) if self.metadata else 0,
            "num_files": len(set(m['source'] for m in self.metadata)) if self.metadata else 0,
            "unique_titles": len(set(m['title'] for m in self.metadata)) if self.metadata else 0,
            "index_type": type(self.faiss_index).__name__ if self.faiss_index else None,
            "index_metadata": self.index_metadata,
        }
        self.logger.info(f"Index stats: {stats}")
        return stats

    def rebuild_index(self) -> None:
        """
        Принудительное перестроение индекса (с backup).
        """
        self.logger.info("Manual index rebuild triggered...")
        notify_admin("Manual HybridRetriever index rebuild triggered by user.")
        self._build_indices()

# код - rag_table_utils.py

import pandas as pd
from pathlib import Path
from rag_file_utils import clean_html_from_cell
import logging

logger = logging.getLogger("rag_table_utils")
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s [%(levelname)s] [rag_table_utils] %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)

def process_table_for_rag(
    file_path: Path,
    columns=None,
    filter_expr=None,
    add_headers=True,
    row_delim="\n"
) -> str:
    ext = file_path.suffix.lower()
    try:
        logger.info(f"Processing table for RAG: {file_path.name}")
        if ext == ".csv":
            df = pd.read_csv(file_path, usecols=columns)
        elif ext in [".xlsx", ".xls", ".xlsm"]:
            df = pd.read_excel(file_path, usecols=columns)
        else:
            logger.error("Not a table file")
            raise ValueError("Not a table file")
        if filter_expr:
            df = df.query(filter_expr)
        for col in df.columns:
            df[col] = df[col].apply(clean_html_from_cell)
        rows = []
        colnames = list(df.columns)
        for idx, row in df.iterrows():
            row_items = [f"{col}: {row[col]}" for col in colnames]
            rows.append(" | ".join(row_items))
        result = ""
        if add_headers:
            header = " | ".join(colnames)
            result = header + row_delim
        result += row_delim.join(rows)
        logger.info(f"Table processed for RAG: {file_path.name}, rows: {len(df)}")
        return result
    except Exception as e:
        logger.error(f"process_table_for_rag error: {e}")
        return f"[Ошибка обработки таблицы для RAG]: {e}"

# код - rag_telegram.py

import requests
import json
from pathlib import Path
from typing import Union, Optional, List
from utils.exceptions import TelegramError
import logging
import time
import html
import traceback

def get_logger(name: str, logfile: Optional[Union[str, Path]] = None, level=logging.INFO) -> logging.Logger:
    logger = logging.getLogger(name)
    if not logger.hasHandlers():
        if logfile:
            fh = logging.FileHandler(logfile, encoding="utf-8")
            fh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
            logger.addHandler(fh)
        sh = logging.StreamHandler()
        sh.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
        logger.addHandler(sh)
        logger.setLevel(level)
    return logger

logger = get_logger("rag_telegram")

def escape_html(text: str) -> str:
    """
    Экранирует HTML-спецсимволы для Telegram (HTML-mode).
    """
    return html.escape(text, quote=False)

class TelegramPublisher:
    """
    Публикация сообщений и файлов в Telegram-канал через Bot API.
    Поддерживает отправку текста, изображений, документов, видео, аудио, предпросмотр ссылок, отложенную публикацию.
    """

    def __init__(
        self,
        bot_token: str,
        channel_id: Union[str, int],
        logger: Optional[logging.Logger] = None,
        max_retries: int = 3,
        retry_delay: float = 3.0,
        enable_preview: bool = True
    ):
        """
        :param bot_token: Токен Telegram-бота
        :param channel_id: ID или username канала (например, @my_channel)
        :param logger: Логгер
        :param max_retries: Количество попыток при ошибках сети/Telegram
        :param retry_delay: Задержка между попытками (сек)
        :param enable_preview: Включить предпросмотр ссылок в постах
        """
        self.bot_token = bot_token
        self.channel_id = channel_id
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.enable_preview = enable_preview
        self.logger = logger or get_logger("rag_telegram")

    def _post(self, method: str, data: dict, files: dict = None) -> dict:
        url = f"https://api.telegram.org/bot{self.bot_token}/{method}"
        last_exc = None
        for attempt in range(1, self.max_retries + 1):
            try:
                resp = requests.post(url, data=data, files=files, timeout=20)
                resp.raise_for_status()
                result = resp.json()
                if not result.get("ok"):
                    self.logger.error(f"Telegram API error: {result}")
                    raise TelegramError(f"Telegram API error: {result}")
                return result
            except Exception as e:
                last_exc = e
                tb = traceback.format_exc()
                self.logger.warning(f"Telegram API request failed (attempt {attempt}): {e}\n{tb}")
                time.sleep(self.retry_delay)
        self.logger.error(f"Telegram API request failed after {self.max_retries} attempts: {last_exc}")
        raise TelegramError(f"Telegram API request failed after {self.max_retries} attempts: {last_exc}") from last_exc

    def send_text(
        self,
        text: str,
        parse_mode: str = "HTML",
        disable_preview: Optional[bool] = None,
        reply_to_message_id: Optional[int] = None,
        silent: bool = False,
        html_escape: bool = True
    ) -> Optional[int]:
        """
        Отправка текстового сообщения в канал.
        :param html_escape: экранировать HTML-спецсимволы (True по умолчанию)
        :return: message_id отправленного сообщения или None при ошибке
        """
        if html_escape and parse_mode == "HTML":
            text = escape_html(text)
        data = {
            "chat_id": self.channel_id,
            "text": text,
            "parse_mode": parse_mode,
            "disable_web_page_preview": not (disable_preview if disable_preview is not None else self.enable_preview),
            "disable_notification": silent,
        }
        if reply_to_message_id:
            data["reply_to_message_id"] = reply_to_message_id
        try:
            resp = self._post("sendMessage", data)
            msg_id = resp.get("result", {}).get("message_id")
            self.logger.info(f"Message posted to Telegram (id={msg_id})")
            return msg_id
        except Exception as e:
            self.logger.error(f"Failed to send text message: {e}\n{traceback.format_exc()}")
            return None

    def send_photo(
        self,
        photo: Union[str, Path],
        caption: Optional[str] = None,
        parse_mode: str = "HTML",
        disable_preview: Optional[bool] = None,
        reply_to_message_id: Optional[int] = None,
        silent: bool = False,
        html_escape: bool = True
    ) -> Optional[int]:
        """
        Отправка фото с подписью.
        :param photo: путь к файлу или URL
        :param html_escape: экранировать HTML в подписи
        :return: message_id или None
        """
        data = {
            "chat_id": self.channel_id,
            "parse_mode": parse_mode,
            "disable_notification": silent,
        }
        if caption:
            data["caption"] = escape_html(caption) if html_escape and parse_mode == "HTML" else caption
        if reply_to_message_id:
            data["reply_to_message_id"] = reply_to_message_id
        files = {}
        file_handle = None
        try:
            if isinstance(photo, (str, Path)) and Path(photo).exists():
                file_handle = open(photo, "rb")
                files["photo"] = file_handle
            else:
                data["photo"] = str(photo)
            resp = self._post("sendPhoto", data, files)
            msg_id = resp.get("result", {}).get("message_id")
            self.logger.info(f"Photo posted to Telegram (id={msg_id})")
            return msg_id
        except Exception as e:
            self.logger.error(f"Failed to send photo: {e}\n{traceback.format_exc()}")
            return None
        finally:
            if file_handle:
                file_handle.close()

    def send_video(
        self,
        video: Union[str, Path],
        caption: Optional[str] = None,
        parse_mode: str = "HTML",
        reply_to_message_id: Optional[int] = None,
        silent: bool = False,
        html_escape: bool = True
    ) -> Optional[int]:
        """
        Отправка видеофайла.
        """
        data = {
            "chat_id": self.channel_id,
            "parse_mode": parse_mode,
            "disable_notification": silent,
        }
        if caption:
            data["caption"] = escape_html(caption) if html_escape and parse_mode == "HTML" else caption
        if reply_to_message_id:
            data["reply_to_message_id"] = reply_to_message_id
        files = {}
        file_handle = None
        try:
            if isinstance(video, (str, Path)) and Path(video).exists():
                file_handle = open(video, "rb")
                files["video"] = file_handle
            else:
                data["video"] = str(video)
            resp = self._post("sendVideo", data, files)
            msg_id = resp.get("result", {}).get("message_id")
            self.logger.info(f"Video posted to Telegram (id={msg_id})")
            return msg_id
        except Exception as e:
            self.logger.error(f"Failed to send video: {e}\n{traceback.format_exc()}")
            return None
        finally:
            if file_handle:
                file_handle.close()

    def send_audio(
        self,
        audio: Union[str, Path],
        caption: Optional[str] = None,
        parse_mode: str = "HTML",
        reply_to_message_id: Optional[int] = None,
        silent: bool = False,
        html_escape: bool = True
    ) -> Optional[int]:
        """
        Отправка аудиофайла.
        """
        data = {
            "chat_id": self.channel_id,
            "parse_mode": parse_mode,
            "disable_notification": silent,
        }
        if caption:
            data["caption"] = escape_html(caption) if html_escape and parse_mode == "HTML" else caption
        if reply_to_message_id:
            data["reply_to_message_id"] = reply_to_message_id
        files = {}
        file_handle = None
        try:
            if isinstance(audio, (str, Path)) and Path(audio).exists():
                file_handle = open(audio, "rb")
                files["audio"] = file_handle
            else:
                data["audio"] = str(audio)
            resp = self._post("sendAudio", data, files)
            msg_id = resp.get("result", {}).get("message_id")
            self.logger.info(f"Audio posted to Telegram (id={msg_id})")
            return msg_id
        except Exception as e:
            self.logger.error(f"Failed to send audio: {e}\n{traceback.format_exc()}")
            return None
        finally:
            if file_handle:
                file_handle.close()

    def send_document(
        self,
        document: Union[str, Path],
        caption: Optional[str] = None,
        parse_mode: str = "HTML",
        reply_to_message_id: Optional[int] = None,
        silent: bool = False,
        html_escape: bool = True
    ) -> Optional[int]:
        """
        Отправка файла-документа.
        :param document: путь к файлу или URL
        :param html_escape: экранировать HTML в подписи
        :return: message_id или None
        """
        data = {
            "chat_id": self.channel_id,
            "parse_mode": parse_mode,
            "disable_notification": silent,
        }
        if caption:
            data["caption"] = escape_html(caption) if html_escape and parse_mode == "HTML" else caption
        if reply_to_message_id:
            data["reply_to_message_id"] = reply_to_message_id
        files = {}
        file_handle = None
        try:
            if isinstance(document, (str, Path)) and Path(document).exists():
                file_handle = open(document, "rb")
                files["document"] = file_handle
            else:
                data["document"] = str(document)
            resp = self._post("sendDocument", data, files)
            msg_id = resp.get("result", {}).get("message_id")
            self.logger.info(f"Document posted to Telegram (id={msg_id})")
            return msg_id
        except Exception as e:
            self.logger.error(f"Failed to send document: {e}\n{traceback.format_exc()}")
            return None
        finally:
            if file_handle:
                file_handle.close()

    def send_media_group(
        self,
        media: List[dict]
    ) -> Optional[List[int]]:
        """
        Отправка набора медиа (фото/видео) в одном сообщении.
        :param media: список dict с типом ('photo'/'video'), media (file_id/url), caption (optional)
        :return: список message_id или None
        """
        data = {
            "chat_id": self.channel_id,
            "media": json.dumps(media, ensure_ascii=False),
        }
        try:
            resp = self._post("sendMediaGroup", data)
            results = resp.get("result", [])
            msg_ids = [msg.get("message_id") for msg in results if "message_id" in msg]
            self.logger.info(f"Media group posted to Telegram (messages={msg_ids})")
            return msg_ids
        except Exception as e:
            self.logger.error(f"Failed to send media group: {e}\n{traceback.format_exc()}")
            return None

    def check_connection(self) -> bool:
        """
        Проверка связи с Telegram Bot API (getMe).
        """
        url = f"https://api.telegram.org/bot{self.bot_token}/getMe"
        try:
            resp = requests.get(url, timeout=10)
            resp.raise_for_status()
            data = resp.json()
            if data.get("ok"):
                self.logger.info("Telegram bot connection OK")
                return True
            else:
                self.logger.error("Telegram bot connection failed")
                return False
        except Exception as e:
            self.logger.error(f"Telegram bot connection error: {e}\n{traceback.format_exc()}")
            return False

    def delayed_post(
        self,
        text: str,
        delay_sec: float,
        **kwargs
    ) -> Optional[int]:
        """
        Отправка сообщения с задержкой.
        """
        self.logger.info(f"Delaying message post for {delay_sec} seconds...")
        time.sleep(delay_sec)
        return self.send_text(text, **kwargs)

# код - search_utils.py

import os
import json
import logging
from typing import List, Dict, Optional, Union, Tuple
from dataclasses import dataclass
from pathlib import Path
import uuid
from datetime import datetime

import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Класс для представления документа"""
    id: str
    content: str
    metadata: Dict = None
    timestamp: str = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.timestamp is None:
            self.timestamp = datetime.now().isoformat()

@dataclass
class QueryResult:
    """Класс для результатов поиска"""
    document_id: str
    content: str
    similarity_score: float
    metadata: Dict
    embedding: List[float] = None

class DocumentProcessor:
    """Класс для предобработки документов"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        # Можно подключить rag_text_utils.clean_html, strip_non_printable и т.д.
        return ' '.join(text.strip().split())
    
     @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50, max_chunks: Optional[int] = None) -> List[str]:
        words = text.split()
        if not words:
            return []
        chunk_size = max(chunk_size, 1)
        overlap = min(overlap, chunk_size-1) if chunk_size > 1 else 0
        step = max(chunk_size - overlap, 1)
        chunks = []
        for i in range(0, len(words), step):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
            if max_chunks and len(chunks) >= max_chunks:
                break
        return chunks
    
    @classmethod
    def process_document(cls, content: str, chunk_size: int = 500, overlap: int = 50, max_chunks: Optional[int] = None) -> List[str]:
        cleaned = cls.clean_text(content)
        return cls.chunk_text(cleaned, chunk_size=chunk_size, overlap=overlap, max_chunks=max_chunks)

class EmbeddingManager:
    """Управление эмбеддингами"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)
        logger.info(f"Loaded embedding model: {model_name}")
    
    def encode(self, texts: Union[str, List[str]], **kwargs) -> np.ndarray:
         try:
            if isinstance(texts, str):
                texts = [texts]
            embeddings = self.model.encode(texts, **kwargs)
            return np.asarray(embeddings)
        except Exception as e:
            logger.error(f"Ошибка создания эмбеддингов: {e}")
            return np.zeros((len(texts), self.model.get_sentence_embedding_dimension()))
    
    def compute_similarity(self, embedding1: np.ndarray, embedding2: np.ndarray) -> float:
        try:
            if embedding1.ndim == 1:
                embedding1 = embedding1.reshape(1, -1)
            if embedding2.ndim == 1:
                embedding2 = embedding2.reshape(1, -1)
            return float(cosine_similarity(embedding1, embedding2)[0][0])
        except Exception as e:
            logger.error(f"Ошибка вычисления similarity: {e}")
            return 0.0

class AdvancedRAGPipeline:
    """Продвинутый RAG pipeline"""
    
    def __init__(
        self,
        collection_name: str = "advanced_rag",
        persist_directory: Optional[str] = None,
        embedding_model: str = "all-MiniLM-L6-v2"
    ):
        self.collection_name = collection_name
        self.persist_directory = persist_directory
        
        # Инициализация ChromaDB
        self._init_chromadb()
        
        # Инициализация компонентов
        self.embedding_manager = EmbeddingManager(embedding_model)
        self.processor = DocumentProcessor()
        
        logger.info("RAG Pipeline initialized successfully")
    
    def _init_chromadb(self):
        """Инициализация ChromaDB"""
        settings = Settings()
        if self.persist_directory:
            settings.persist_directory = self.persist_directory
            Path(self.persist_directory).mkdir(parents=True, exist_ok=True)
        
        self.client = chromadb.Client(settings)
        
        # Создание или получение коллекции
        try:
            self.collection = self.client.get_collection(self.collection_name)
            logger.info(f"Loaded existing collection: {self.collection_name}")
        except:
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"description": "Advanced RAG Pipeline Collection"}
            )
            logger.info(f"Created new collection: {self.collection_name}")
    
    def add_document(
        self,
        content: str,
        doc_id: Optional[str] = None,
        metadata: Optional[Dict] = None,
        chunk_size: int = 500,
        overlap: int = 50,
        max_chunks: int = 1000,
        max_doc_len: int = 100_000
    ) -> List[str]:
        if not content or not isinstance(content, str):
            logger.warning("Пустой или невалидный контент, документ не будет добавлен")
            return []
        if len(content) > max_doc_len:
            logger.warning(f"Документ слишком большой ({len(content)}), будет усечён до {max_doc_len}")
            content = content[:max_doc_len]
        
        # Обработка документа
        chunks = self.processor.process_document(content, chunk_size)
        
        # Создание эмбеддингов
        embeddings = self.embedding_manager.encode(chunks)
        
        # Подготовка данных для ChromaDB
        chunk_ids = []
        chunk_documents = []
        chunk_embeddings = []
        chunk_metadata = []
        
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{doc_id}_chunk_{i}"
            chunk_ids.append(chunk_id)
            chunk_documents.append(chunk)
            chunk_embeddings.append(embedding.tolist())
            
            chunk_meta = {
                **metadata,
                "parent_doc_id": doc_id,
                "chunk_index": i,
                "chunk_count": len(chunks),
                "timestamp": datetime.now().isoformat()
            }
            chunk_metadata.append(chunk_meta)
        
        # Добавление в коллекцию
        self.collection.add(
            ids=chunk_ids,
            documents=chunk_documents,
            embeddings=chunk_embeddings,
            metadatas=chunk_metadata
        )
        
        logger.info(f"Added document {doc_id} with {len(chunks)} chunks")
        return chunk_ids
    
    def add_documents_batch(
        self,
        documents: List[Document],
        chunk_size: int = 500,
        batch_size: int = 100
    ) -> List[str]:
        """Пакетное добавление документов"""
        all_chunk_ids = []
        
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            batch_ids = []
            
            for doc in batch:
                chunk_ids = self.add_document(
                    content=doc.content,
                    doc_id=doc.id,
                    metadata=doc.metadata,
                    chunk_size=chunk_size
                )
                batch_ids.extend(chunk_ids)
            
            all_chunk_ids.extend(batch_ids)
            logger.info(f"Processed batch {i//batch_size + 1}: {len(batch)} documents")
        
        return all_chunk_ids
    
    def search(
        self,
        query: str,
        n_results: int = 5,
        filter_metadata: Optional[Dict] = None,
        min_similarity: float = 0.0
    ) -> List[QueryResult]:
        """Поиск документов"""
        # Создание эмбеддинга запроса
        query_embedding = self.embedding_manager.encode([query])[0]
        
        # Поиск в ChromaDB
        where_clause = filter_metadata if filter_metadata else None
        
        results = self.collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=n_results,
            where=where_clause
        )
        
        # Обработка результатов
        query_results = []
        
        for i in range(len(results['ids'][0])):
            doc_id = results['ids'][0][i]
            content = results['documents'][0][i]
            metadata = results['metadatas'][0][i]
            distance = results['distances'][0][i] if 'distances' in results else None
            
            # Преобразование distance в similarity score
            similarity_score = 1 - distance if distance is not None else 0.0
            
            if similarity_score >= min_similarity:
                query_result = QueryResult(
                    document_id=doc_id,
                    content=content,
                    similarity_score=similarity_score,
                    metadata=metadata
                )
                query_results.append(query_result)
        
        logger.info(f"Found {len(query_results)} results for query: '{query[:50]}...'")
        return query_results
    
    def delete_document(self, doc_id: str) -> bool:
        """Удаление документа и всех его чанков"""
        try:
            # Находим все чанки документа
            results = self.collection.get(
                where={"parent_doc_id": doc_id}
            )
            
            if results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.info(f"Deleted document {doc_id} and {len(results['ids'])} chunks")
                return True
            else:
                logger.warning(f"Document {doc_id} not found")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")
            return False
    
    def update_document(
        self,
        doc_id: str,
        new_content: str,
        new_metadata: Optional[Dict] = None,
        chunk_size: int = 500
    ) -> List[str]:
        """Обновление документа"""
        # Удаляем старый документ
        self.delete_document(doc_id)
        
        # Добавляем новый
        return self.add_document(
            content=new_content,
            doc_id=doc_id,
            metadata=new_metadata,
            chunk_size=chunk_size
        )
    
    def get_collection_stats(self) -> Dict:
        """Статистика коллекции"""
        try:
            count = self.collection.count()
            
            # Получаем примеры метаданных
            sample = self.collection.peek(limit=10)
            
            stats = {
                "total_chunks": count,
                "collection_name": self.collection_name,
                "sample_metadata": sample.get('metadatas', [])[:3] if sample else []
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}
    
    def export_collection(self, filepath: str, chunk_limit: int = 100_000) -> bool:
        """Экспорт коллекции в JSON"""
        try:
            all_data = self.collection.get(limit=chunk_limit)
            
            export_data = {
                "collection_name": self.collection_name,
                "export_timestamp": datetime.now().isoformat(),
                "data": {
                    "ids": all_data.get('ids', []),
                    "documents": all_data.get('documents', []),
                    "metadatas": all_data.get('metadatas', []),
                    "embeddings": all_data.get('embeddings', [])
                }
            }
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"Collection exported to {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error exporting collection: {e}")
            return False
    
    def semantic_search_with_reranking(
        self,
        query: str,
        n_results: int = 10,
        rerank_top_k: int = 5
    ) -> List[QueryResult]:
        initial_results = self.search(query, n_results=n_results)
        if not initial_results:
            return []
        query_embedding = self.embedding_manager.encode([query])[0]
        # Кэшируем эмбеддинги документов (чтобы не считать повторно)
        doc_texts = [r.content for r in initial_results]
        doc_embeddings = self.embedding_manager.encode(doc_texts)
        reranked_results = []
        for i, result in enumerate(initial_results):
            similarity = self.embedding_manager.compute_similarity(query_embedding, doc_embeddings[i])
            result.similarity_score = similarity
            reranked_results.append(result)
        reranked_results.sort(key=lambda x: x.similarity_score, reverse=True)
        return reranked_results[:rerank_top_k]

# Пример использования
if __name__ == "__main__":
    # Инициализация RAG pipeline
    rag = AdvancedRAGPipeline(
        collection_name="advanced_rag_demo",
        persist_directory="./rag_storage"
    )
    
    # Тестовые документы
    test_documents = [
        Document(
            id="doc1",
            content="Векторные базы данных используются для хранения и поиска эмбеддингов. ChromaDB - одна из популярных векторных баз данных для Python разработки.",
            metadata={"category": "database", "language": "python"}
        ),
        Document(
            id="doc2", 
            content="RAG (Retrieval-Augmented Generation) pipeline объединяет поиск релевантной информации с генерацией ответов. Это мощный подход для создания AI-ассистентов.",
            metadata={"category": "ai", "topic": "rag"}
        ),
        Document(
            id="doc3",
            content="SentenceTransformers предоставляет предобученные модели для создания качественных эмбеддингов текста. Модель all-MiniLM-L6-v2 подходит для большинства задач.",
            metadata={"category": "ml", "library": "sentence_transformers"}
        )
    ]
    
    # Добавление документов
    print("Добавление документов...")
    rag.add_documents_batch(test_documents)
    
    # Статистика
    print("\nСтатистика коллекции:")
    stats = rag.get_collection_stats()
    print(json.dumps(stats, indent=2, ensure_ascii=False))
    
    # Поиск
    print("\nПоиск документов:")
    query = "как использовать векторные базы данных в Python"
    results = rag.search(query, n_results=3)
    
    for i, result in enumerate(results):
        print(f"\nРезультат {i+1}:")
        print(f"ID: {result.document_id}")
        print(f"Схожесть: {result.similarity_score:.3f}")
        print(f"Контент: {result.content[:100]}...")
        print(f"Метаданные: {result.metadata}")
    
    # Семантический поиск с переранжированием
    print("\nПоиск с переранжированием:")
    reranked = rag.semantic_search_with_reranking(query, n_results=5, rerank_top_k=2)
    
    for i, result in enumerate(reranked):
        print(f"\nПереранжированный результат {i+1}:")
        print(f"Схожесть: {result.similarity_score:.3f}")
        print(f"Контент: {result.content[:100]}...")

# код - image_utils.py

import random
from pathlib import Path
from typing import Optional, Tuple, List, Set, Union
from PIL import Image, UnidentifiedImageError
import logging
import mimetypes
from enum import Enum, auto

logger = logging.getLogger("image_utils")
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] [image_utils] %(message)s'))
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

SUPPORTED_IMAGE_EXTS: Set[str] = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
SUPPORTED_VIDEO_EXTS: Set[str] = {".mp4", ".mov", ".mkv"}
SUPPORTED_DOC_EXTS: Set[str] = {".pdf", ".docx", ".doc", ".txt", ".csv", ".xlsx"}
SUPPORTED_AUDIO_EXTS: Set[str] = {".mp3", ".wav", ".ogg"}
SUPPORTED_MEDIA_EXTS: Set[str] = SUPPORTED_IMAGE_EXTS | SUPPORTED_VIDEO_EXTS | SUPPORTED_DOC_EXTS | SUPPORTED_AUDIO_EXTS

MAX_IMAGE_SIZE: Tuple[int, int] = (1280, 1280)
MAX_FILE_SIZE_MB: int = 50

class MediaValidationStatus(Enum):
    OK = auto()
    NOT_FOUND = auto()
    SYMLINK = auto()
    OUT_OF_DIR = auto()
    NOT_SUPPORTED = auto()
    TOO_LARGE = auto()
    CANNOT_OPEN = auto()
    UNKNOWN = auto()

class MediaValidationResult:
    def __init__(self, valid: bool, status: MediaValidationStatus, path: Optional[Path], message: str = ""):
        self.valid = valid
        self.status = status
        self.path = path
        self.message = message

    def __bool__(self):
        return self.valid

    def __repr__(self):
        return f"<MediaValidationResult status={self.status} path={self.path} msg={self.message}>"

def is_safe_media_path(path: Path, media_dir: Path) -> bool:
    """
    Проверяет, что path лежит строго внутри media_dir.
    Защита от directory traversal.
    """
    try:
        # Python >=3.9
        return path.resolve().is_relative_to(media_dir.resolve())
    except AttributeError:  # Python <3.9
        try:
            return str(path.resolve().as_posix()).startswith(media_dir.resolve().as_posix() + "/")
        except Exception as e:
            logger.error(f"Ошибка проверки пути: {e}")
            return False
    except Exception as e:
        logger.error(f"Ошибка проверки пути: {e}")
        return False

def pick_random_media_file(media_dir: Path, allowed_exts: Optional[Set[str]] = None, max_retries: int = 5) -> Optional[Path]:
    """
    Случайно выбирает валидный файл из media_dir (рекурсивно) с поддерживаемым расширением.
    Делает max_retries попыток выбрать валидный файл.
    """
    allowed_exts = allowed_exts or SUPPORTED_MEDIA_EXTS
    files = [f for f in media_dir.rglob("*") if f.is_file() and f.suffix.lower() in allowed_exts]
    if not files:
        logger.warning(f"Нет файлов с поддерживаемыми расширениями в {media_dir}")
        return None
    attempt = 0
    while attempt < max_retries and files:
        file = random.choice(files)
        result = validate_media_file(file, media_dir)
        if result.valid:
            return file
        else:
            logger.info(f"Пропущен невалидный файл при попытке выбора: {file} ({result.status})")
            files.remove(file)
            attempt += 1
    logger.warning("Не удалось выбрать валидный файл после повторных попыток.")
    return None

def validate_media_file(path: Path, media_dir: Path = Path("media")) -> MediaValidationResult:
    """
    Проверяет валидность файла:
    - В media_dir
    - Поддерживаемое расширение
    - Не превышает лимит размера
    - Существует
    - Не symlink
    """
    if not path.exists():
        logger.error(f"Файл не найден: {path}")
        return MediaValidationResult(False, MediaValidationStatus.NOT_FOUND, path, "Файл не найден")
    if path.is_symlink():
        logger.error(f"Файл является symlink: {path}")
        return MediaValidationResult(False, MediaValidationStatus.SYMLINK, path, "Файл является symlink")
    if not is_safe_media_path(path, media_dir):
        logger.error(f"Файл вне папки media: {path}")
        return MediaValidationResult(False, MediaValidationStatus.OUT_OF_DIR, path, "Файл вне папки media")
    if path.suffix.lower() not in SUPPORTED_MEDIA_EXTS:
        logger.error(f"Неподдерживаемый формат: {path.suffix} ({path})")
        return MediaValidationResult(False, MediaValidationStatus.NOT_SUPPORTED, path, f"Неподдерживаемый формат: {path.suffix}")
    if path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        logger.error(f"Файл слишком большой (> {MAX_FILE_SIZE_MB} МБ): {path}")
        return MediaValidationResult(False, MediaValidationStatus.TOO_LARGE, path, f"Файл слишком большой (>{MAX_FILE_SIZE_MB} МБ)")
    mime, _ = mimetypes.guess_type(str(path))
    if mime is None:
        logger.warning(f"Не удалось определить MIME-тип: {path}")
    return MediaValidationResult(True, MediaValidationStatus.OK, path, "OK")

def get_media_type(path: Path) -> str:
    """
    Определяет тип файла по расширению.
    """
    ext = path.suffix.lower()
    if ext in SUPPORTED_IMAGE_EXTS:
        return "image"
    elif ext in SUPPORTED_VIDEO_EXTS:
        return "video"
    elif ext in SUPPORTED_DOC_EXTS:
        return "document"
    elif ext in SUPPORTED_AUDIO_EXTS:
        return "audio"
    return "unknown"

def process_image(path: Path, output_dir: Optional[Path] = None, max_size: Tuple[int, int] = MAX_IMAGE_SIZE) -> Optional[Path]:
    """
    Уменьшает изображение до max_size, если требуется. Возвращает путь к новому файлу.
    """
    try:
        with Image.open(path) as img:
            if img.size[0] <= max_size[0] and img.size[1] <= max_size[1]:
                logger.info(f"Изображение уже в допустимом размере: {path}")
                return path
            img.thumbnail(max_size, Image.LANCZOS)
            out_dir = output_dir or path.parent
            out_path = out_dir / f"{path.stem}_resized{path.suffix}"
            img.save(out_path)
            logger.info(f"Изображение уменьшено и сохранено в: {out_path}")
            return out_path
    except UnidentifiedImageError:
        logger.error(f"Не удалось открыть изображение: {path}")
        return None
    except Exception as e:
        logger.error(f"Ошибка обработки изображения {path}: {e}")
        return None

def get_all_media_files(media_dir: Path, allowed_exts: Optional[Set[str]] = None) -> List[Path]:
    """
    Список всех файлов в media_dir c поддерживаемыми расширениями.
    """
    allowed_exts = allowed_exts or SUPPORTED_MEDIA_EXTS
    return [f for f in media_dir.rglob("*") if f.is_file() and f.suffix.lower() in allowed_exts]

def prepare_media_for_post(media_dir: Path = Path("media")) -> Optional[Path]:
    """
    Выбирает и валидирует случайный медиа-файл. При необходимости уменьшает изображение.
    """
    file = pick_random_media_file(media_dir)
    if not file:
        logger.warning("Не найден ни один подходящий медиа-файл для публикации.")
        return None
    result = validate_media_file(file, media_dir)
    if not result.valid:
        logger.error(f"Медиа-файл не прошёл валидацию: {file}. Причина: {result.status} ({result.message})")
        return None
    media_type = get_media_type(file)
    if media_type == "image":
        try:
            with Image.open(file) as img:
                if img.size[0] > MAX_IMAGE_SIZE[0] or img.size[1] > MAX_IMAGE_SIZE[1]:
                    resized = process_image(file)
                    if resized is not None:
                        return resized
                return file
        except Exception as e:
            logger.error(f"Ошибка открытия изображения: {e}")
            return None
    return file

# код - rag_text_utils.py

from pathlib import Path
from typing import List, Union
from logs import get_logger

logger = get_logger("rag_text_utils")

COMMON_ENCODINGS = ["utf-8", "cp1251", "windows-1251", "latin-1", "iso-8859-1"]
MAX_FILE_SIZE_MB = 50

def _smart_read_text(path: Path) -> str:
    """
    Пробует прочитать текстовый файл с помощью популярных кодировок.
    Возвращает содержимое файла или выбрасывает UnicodeDecodeError, если не удалось.
    """
    if path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        logger.error(f"Файл слишком большой (> {MAX_FILE_SIZE_MB} МБ): {path}")
        raise IOError(f"File too large: {path}")
    for encoding in COMMON_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except Exception as e:
            logger.debug(f"Проблема с кодировкой {encoding} для {path}: {e}")
    logger.error(f"Не удалось прочитать файл {path} в поддерживаемых кодировках: {COMMON_ENCODINGS}")
    raise UnicodeDecodeError("all", b'', 0, 1, f"Failed to read {path} with encodings: {COMMON_ENCODINGS}")

def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 0,
    max_chunks: int = None
) -> List[str]:
    """
    Делит строку на чанки по словам с заданным размером и overlap.
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size должен быть положительным")
    if overlap < 0:
        raise ValueError("overlap не может быть отрицательным")
    words = text.split()
    if not words:
        return []
    chunks = []
    step = max(chunk_size - overlap, 1)
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i+chunk_size])
        if chunk.strip():  # отбрасываем пустые чанки
            chunks.append(chunk)
        if max_chunks is not None and len(chunks) >= max_chunks:
            logger.info(f"Обрезано по max_chunks={max_chunks}")
            break
    return chunks

def process_text_file_for_rag(
    file_path: Path,
    chunk_size: int = 1000,
    overlap: int = 0,
    max_chunks: int = None,
    raise_on_error: bool = False
) -> List[str]:
    """
    Читает текстовый файл и делит его на чанки для RAG.
    """
    try:
        text = _smart_read_text(file_path)
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap, max_chunks=max_chunks)
        logger.info(f"Text file processed for RAG: {file_path.name}, chunks: {len(chunks)}")
        return chunks
    except Exception as e:
        logger.error(f"process_text_file_for_rag error: {e}")
        if raise_on_error:
            raise
        return []

def process_text_for_rag(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 0,
    max_chunks: int = None
) -> List[str]:
    """
    Делит произвольную строку на чанки для RAG.
    """
    try:
        chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap, max_chunks=max_chunks)
        logger.info(f"Arbitrary text processed for RAG, chunks: {len(chunks)}")
        return chunks
    except Exception as e:
        logger.error(f"process_text_for_rag error: {e}")
        return []

# Возможное расширение: чанкинг по предложениям и абзацам, асинхронное чтение, поддержка stream-обработки.

# код - rag_prompt_utils.py

from pathlib import Path
from typing import Optional, Union, Dict, Any
from logs import get_logger
import random
import re

logger = get_logger("rag_prompt_utils")

class PromptTemplateCache:
    _cache: Dict[str, str] = {}
    @classmethod
    def get(cls, path: Path) -> Optional[str]:
        key = str(path.resolve())
        if key in cls._cache:
            return cls._cache[key]
        if not path.exists():
            logger.warning(f"Шаблон промпта не найден: {path}")
            return None
        try:
            text = path.read_text(encoding="utf-8")
            cls._cache[key] = text
            return text
        except Exception as e:
            logger.error(f"Ошибка чтения шаблона: {path}: {e}")
            return None

def safe_format(template: str, variables: Dict[str, Any]) -> str:
    def repl(match):
        var = match.group(1)
        return str(variables.get(var, f"{{{var}}}"))
    return re.sub(r"\{(\w+)\}", repl, template)

def get_prompt_parts(
    data_dir: Union[str, Path],
    topic: str,
    context: str,
    uploadfile: Optional[Union[str, Path]] = None,
    file1: Optional[Union[str, Path]] = None,
    file2: Optional[Union[str, Path]] = None,
    extra_vars: Optional[Dict[str, Any]] = None,
    max_context_len_upload: int = 1024,
    max_context_len_no_upload: int = 4096,
    max_prompt_len: int = 8192,
) -> str:
    """
    Составляет промпт для LLM на основе шаблонов и переданных параметров.
    Условия:
      - {TOPIC}: каждая строка из topics.txt, подставляется отдельно
      - {CONTEXT}: весь материал из RAG+интернета, лимит 4096 или 1024 если есть {UPLOADFILE}
      - {UPLOADFILE}: имя рандомного файла из media или статус ошибки
    """
    data_dir = Path(data_dir)
    if file1 is not None:
        file1 = Path(file1)
    if file2 is not None:
        file2 = Path(file2)
    if uploadfile is not None:
        uploadfile_path = Path(uploadfile)
    else:
        uploadfile_path = None

    # Проверка существования data_dir
    if not data_dir.exists() or not data_dir.is_dir():
        logger.error(f"data_dir '{data_dir}' не существует или не является директорией")
        return f"{topic}\n\n{context[:max_context_len_no_upload]}"

    def read_template(path: Path) -> Optional[str]:
        return PromptTemplateCache.get(path)

    prompt1_dir = data_dir / "prompt_1"
    prompt2_dir = data_dir / "prompt_2"
    template = None

    # Детерминированный шаблон
    if file1 is not None and file2 is not None and file1.exists() and file2.exists():
        logger.info(f"Детерминированный шаблон: {file1.name} + {file2.name}")
        txt1 = read_template(file1)
        txt2 = read_template(file2)
        if txt1 is not None and txt2 is not None:
            template = txt1 + "\n" + txt2
    # Случайные шаблоны
    elif prompt1_dir.exists() and prompt2_dir.exists():
        prompt1_files = list(prompt1_dir.glob("*.txt"))
        prompt2_files = list(prompt2_dir.glob("*.txt"))
        if prompt1_files and prompt2_files:
            f1 = random.choice(prompt1_files)
            f2 = random.choice(prompt2_files)
            logger.info(f"Случайный шаблон: {f1.name} + {f2.name}")
            txt1 = read_template(f1)
            txt2 = read_template(f2)
            if txt1 is not None and txt2 is not None:
                template = txt1 + "\n" + txt2

    # Fallback
    if template is None:
        prompt_file = data_dir / "prompt.txt"
        if prompt_file.exists():
            logger.warning("Fallback на prompt.txt")
            template = read_template(prompt_file)
        else:
            logger.warning("Fallback на plain topic + context")
            return f"{topic}\n\n{context[:max_context_len_no_upload]}"

    if template is None or not template.strip():
        logger.error("Шаблон пустой или не удалось прочитать, возврат plain topic + context")
        return f"{topic}\n\n{context[:max_context_len_no_upload]}"

    # Определяем лимит context
    has_uploadfile = "{UPLOADFILE}" in template
    if has_uploadfile:
        context = context[:max_context_len_upload]
    else:
        context = context[:max_context_len_no_upload]

    # Формируем переменные для шаблона
    variables = {
        "TOPIC": topic,
        "CONTEXT": context,
    }
    if extra_vars:
        variables.update(extra_vars)
    # uploadfile logic
    if has_uploadfile:
        if uploadfile_path is not None:
            try:
                if uploadfile_path.exists():
                    variables["UPLOADFILE"] = uploadfile_path.name
                else:
                    variables["UPLOADFILE"] = f"[Файл не найден: {uploadfile_path.name}]"
            except Exception as e:
                variables["UPLOADFILE"] = "[Ошибка с файлом]"
                logger.error(f"Ошибка обработки uploadfile: {e}")
        else:
            variables["UPLOADFILE"] = "[Файл не передан]"

    prompt_out = safe_format(template, variables).strip()
    if len(prompt_out) > max_prompt_len:
        logger.warning(f"Промпт превышает лимит {max_prompt_len}, будет обрезан.")
        prompt_out = prompt_out[:max_prompt_len-10] + "..."
    return prompt_out


# код - RAG_Pipeline_Extensions_Utils.py

import os
import json
import csv
import requests
from typing import List, Dict, Optional
from pathlib import Path
import asyncio
import aiohttp
from dataclasses import dataclass
from utils.path_utils import validate_path
from utils.exceptions import ProcessingError, FileOperationError
import time
import hashlib

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from advanced_rag_pipeline import Document, AdvancedRAGPipeline

class DataIngestionManager:
    """Менеджер для загрузки данных из различных источников"""

    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.allowed_data_dir = Path("./data").resolve()

    def load_from_text_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".txt"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения файла {filepath}: {e}") from e

    def load_from_pdf(self, filepath: str) -> str:
        if not PDF_AVAILABLE:
            raise ProcessingError("PyPDF2 не установлен. Установите: pip install PyPDF2")
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".pdf"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            text = ""
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения PDF {filepath}: {e}") from e

    def load_from_docx(self, filepath: str) -> str:
        if not DOCX_AVAILABLE:
            raise ProcessingError("python-docx не установлен. Установите: pip install python-docx")
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".docx"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения DOCX {filepath}: {e}") from e

    def load_from_csv(self, filepath: str, text_columns: List[str]) -> List[Dict]:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".csv"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            documents = []
            with open(path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    content_parts = [str(row[col]) for col in text_columns if col in row and row[col]]
                    if content_parts:
                        content = " ".join(content_parts)
                        metadata = {k: v for k, v in row.items() if k not in text_columns}
                        documents.append({'id': f"csv_row_{i}", 'content': content, 'metadata': metadata})
            return documents
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения CSV {filepath}: {e}") from e

    def load_from_json(self, filepath: str, content_field: str, id_field: Optional[str] = None) -> List[Dict]:
        path = Path(filepath)
        is_valid, reason = validate_path(
            path, allowed_dir=self.allowed_data_dir, allowed_exts={".json"}, max_size_mb=50
        )
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            documents = []
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if content_field in item:
                        doc_id = item.get(id_field, f"json_item_{i}") if id_field else f"json_item_{i}"
                        content = str(item[content_field])
                        metadata = {k: v for k, v in item.items() if k not in [content_field, id_field]}
                        documents.append({'id': doc_id, 'content': content, 'metadata': metadata})
            return documents
        except Exception as e:
            raise FileOperationError(f"Ошибка чтения JSON {filepath}: {e}") from e

    def load_from_url(self, url: str, timeout: int = 30) -> str:
        try:
            response = requests.get(url, timeout=timeout)
            response.raise_for_status()
            if 'text/html' in response.headers.get('content-type', ''):
                if BS4_AVAILABLE:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    for script in soup(["script", "style"]):
                        script.decompose()
                    return soup.get_text()
                else:
                    return response.text
            else:
                return response.text
        except Exception as e:
            raise ProcessingError(f"Ошибка загрузки URL {url}: {e}") from e

    async def load_from_urls_async(self, urls: List[str], timeout: int = 30) -> List[Dict]:
        async def fetch_url(session, url):
            try:
                async with session.get(url, timeout=timeout) as response:
                    content = await response.text()
                    if 'text/html' in response.headers.get('content-type', ''):
                        if BS4_AVAILABLE:
                            soup = BeautifulSoup(content, 'html.parser')
                            for script in soup(["script", "style"]):
                                script.decompose()
                            content = soup.get_text()
                    return {
                        'id': hashlib.md5(url.encode()).hexdigest(),
                        'content': content,
                        'metadata': {'source_url': url, 'status': 'success'}
                    }
            except Exception as e:
                return {
                    'id': hashlib.md5(url.encode()).hexdigest(),
                    'content': "",
                    'metadata': {'source_url': url, 'status': 'error', 'error': str(e)}
                }
        async with aiohttp.ClientSession() as session:
            tasks = [fetch_url(session, url) for url in urls]
            results = await asyncio.gather(*tasks)
            return [r for r in results if r['content']]

class RAGAnalytics:
    """Аналитика и мониторинг RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.query_log = []

    def log_query(self, query: str, results_count: int, processing_time: float):
        self.query_log.append({
            'timestamp': time.time(),
            'query': query,
            'results_count': results_count,
            'processing_time': processing_time
        })

    def get_query_stats(self) -> Dict:
        if not self.query_log:
            return {"message": "Нет данных по запросам"}
        processing_times = [log['processing_time'] for log in self.query_log]
        results_counts = [log['results_count'] for log in self.query_log]
        return {
            'total_queries': len(self.query_log),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'max_processing_time': max(processing_times),
            'min_processing_time': min(processing_times),
            'avg_results_count': sum(results_counts) / len(results_counts),
            'queries_per_hour': len([q for q in self.query_log if time.time() - q['timestamp'] < 3600])
        }

    def analyze_collection_content(self) -> Dict:
        try:
            all_data = self.rag.collection.get()
            documents = all_data.get('documents', [])
            metadatas = all_data.get('metadatas', [])
            if not documents:
                return {"message": "Коллекция пуста"}
            total_docs = len(documents)
            total_chars = sum(len(doc) for doc in documents)
            avg_doc_length = total_chars / total_docs
            categories = {}
            languages = {}
            for metadata in metadatas:
                if metadata:
                    if 'category' in metadata:
                        cat = metadata['category']
                        categories[cat] = categories.get(cat, 0) + 1
                    if 'language' in metadata:
                        lang = metadata['language']
                        languages[lang] = languages.get(lang, 0) + 1
            return {
                'total_documents': total_docs,
                'total_characters': total_chars,
                'average_document_length': avg_doc_length,
                'categories_distribution': categories,
                'languages_distribution': languages,
                'longest_document': max(len(doc) for doc in documents),
                'shortest_document': min(len(doc) for doc in documents)
            }
        except Exception as e:
            return {"error": f"Ошибка анализа: {e}"}

class RAGWebInterface:
    """Простой веб-интерфейс для RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline, analytics: RAGAnalytics):
        self.rag = rag_pipeline
        self.analytics = analytics

    def generate_html_interface(self) -> str:
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>RAG Pipeline Interface</title>
            <meta charset="UTF-8">
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }
                .container { max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 8px; }
                .search-box { width: 70%; padding: 10px; font-size: 16px; border: 1px solid #ddd; border-radius: 4px; }
                .search-btn { padding: 10px 20px; background: #007bff; color: white; border: none; border-radius: 4px; cursor: pointer; }
                .result { margin: 15px 0; padding: 15px; border: 1px solid #ddd; border-radius: 4px; background: #f9f9f9; }
                .score { color: #666; font-size: 14px; }
                .metadata { color: #888; font-size: 12px; margin-top: 5px; }
                .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }
                .stat-card { padding: 15px; background: #e9ecef; border-radius: 4px; text-align: center; }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>🔍 RAG Pipeline Interface</h1>
                <div class="search-section">
                    <input type="text" class="search-box" id="searchInput" placeholder="Введите ваш запрос...">
                    <button class="search-btn" onclick="performSearch()">Поиск</button>
                </div>
                <div id="results"></div>
                <h2>📊 Статистика системы</h2>
                <div class="stats" id="stats">
                    <div class="stat-card">
                        <h3>Запросы</h3>
                        <div id="queryCount">Загрузка...</div>
                    </div>
                    <div class="stat-card">
                        <h3>Среднее время</h3>
                        <div id="avgTime">Загрузка...</div>
                    </div>
                </div>
                <h2>📈 Управление данными</h2>
                <div style="margin: 20px 0;">
                    <button onclick="exportData()" style="margin-right: 10px; padding: 8px 16px; background: #28a745; color: white; border: none; border-radius: 4px;">Экспорт данных</button>
                    <button onclick="clearCollection()" style="padding: 8px 16px; background: #dc3545; color: white; border: none; border-radius: 4px;">Очистить коллекцию</button>
                </div>
            </div>
            <script>
                async function performSearch() {
                    const query = document.getElementById('searchInput').value;
                    if (!query.trim()) return;
                    const resultsDiv = document.getElementById('results');
                    resultsDiv.innerHTML = '<div>Поиск...</div>';
                    setTimeout(() => {
                        resultsDiv.innerHTML = `
                            <h3>Результаты поиска для: "${query}"</h3>
                            <div class="result">
                                <strong>Документ 1</strong>
                                <div class="score">Релевантность: 0.85</div>
                                <p>Пример найденного контента...</p>
                                <div class="metadata">Метаданные: category=ai, timestamp=2024-01-01</div>
                            </div>
                        `;
                    }, 1000);
                }
                function loadStats() {
                    document.getElementById('docCount').textContent = '150';
                    document.getElementById('queryCount').textContent = '45';
                    document.getElementById('avgTime').textContent = '0.3s';
                }
                function exportData() {
                    alert('Функция экспорта будет реализована на сервере');
                }
                function clearCollection() {
                    if (confirm('Вы уверены, что хотите очистить коллекцию?')) {
                        alert('Функция очистки будет реализована на сервере');
                    }
                }
                window.onload = loadStats;
                document.getElementById('searchInput').addEventListener('keypress', function(e) {
                    if (e.key === 'Enter') {
                        performSearch();
                    }
                });
            </script>
        </body>
        </html>
        """
        return html_template

    def save_interface(self, filepath: str = "rag_interface.html"):
        html_content = self.generate_html_interface()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"HTML интерфейс сохранен в {filepath}")

class RAGBenchmarking:
    """Бенчмаркинг и тестирование RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline

    def create_test_dataset(self, size: int = 100) -> List[Document]:
        test_topics = [
            "машинное обучение", "искусственный интеллект", "нейронные сети",
            "векторные базы данных", "обработка естественного языка",
            "компьютерное зрение", "глубокое обучение", "анализ данных",
            "программирование на Python", "веб-разработка"
        ]
        documents = []
        for i in range(size):
            topic = test_topics[i % len(test_topics)]
            content = self._generate_synthetic_content(topic, i)
            doc = Document(
                id=f"test_doc_{i}",
                content=content,
                metadata={
                    "topic": topic,
                    "test_id": i,
                    "synthetic": True
                }
            )
            documents.append(doc)
        return documents

    def _generate_synthetic_content(self, topic: str, doc_id: int) -> str:
        templates = [
            f"{topic} является важной областью современной технологии. Документ номер {doc_id} содержит подробную информацию о применении {topic} в различных сферах.",
            f"В этом документе рассматриваются основные принципы {topic}. Это руководство номер {doc_id} поможет понять ключевые концепции.",
            f"Практическое применение {topic} демонстрируется в примере {doc_id}. Здесь представлены лучшие практики и методы."
        ]
        base_content = templates[doc_id % len(templates)]
        additional = f" Дополнительные детали включают технические аспекты, примеры использования и рекомендации экспертов в области {topic}."
        return base_content + additional

    def benchmark_search_performance(self, queries: List[str], n_runs: int = 10) -> Dict:
        results = {
            'queries': [],
            'avg_time': 0,
            'total_time': 0,
            'fastest_query': None,
            'slowest_query': None
        }
        total_time = 0
        fastest_time = float('inf')
        slowest_time = 0
        for query in queries:
            query_times = []
            for _ in range(n_runs):
                start_time = time.time()
                self.rag.search(query, n_results=5)
                end_time = time.time()
                query_time = end_time - start_time
                query_times.append(query_time)
                total_time += query_time
            avg_query_time = sum(query_times) / len(query_times)
            min_query_time = min(query_times)
            max_query_time = max(query_times)
            if min_query_time < fastest_time:
                fastest_time = min_query_time
                results['fastest_query'] = query
            if max_query_time > slowest_time:
                slowest_time = max_query_time
                results['slowest_query'] = query
            results['queries'].append({
                'query': query,
                'avg_time': avg_query_time,
                'min_time': min_query_time,
                'max_time': max_query_time,
                'runs': n_runs
            })
        results['avg_time'] = total_time / (len(queries) * n_runs)
        results['total_time'] = total_time
        results['fastest_time'] = fastest_time
        results['slowest_time'] = slowest_time
        return results

    def evaluate_retrieval_quality(self, test_queries: List[Dict]) -> Dict:
        metrics = {
            'precision_at_k': [],
            'recall_at_k': [],
            'average_precision': []
        }
        for test_case in test_queries:
            query = test_case['query']
            relevant_ids = set(test_case['relevant_doc_ids'])
            results = self.rag.search(query, n_results=10)
            retrieved_ids = [r.document_id for r in results]
            for k in [1, 3, 5, 10]:
                if k <= len(retrieved_ids):
                    retrieved_k = set(retrieved_ids[:k])
                    precision = len(retrieved_k & relevant_ids) / k
                    recall = len(retrieved_k & relevant_ids) / len(relevant_ids) if relevant_ids else 0
                    metrics['precision_at_k'].append(precision)
                    metrics['recall_at_k'].append(recall)
            ap = self._calculate_average_precision(retrieved_ids, relevant_ids)
            metrics['average_precision'].append(ap)
        final_metrics = {
            'mean_precision_at_k': sum(metrics['precision_at_k']) / len(metrics['precision_at_k']) if metrics['precision_at_k'] else 0,
            'mean_recall_at_k': sum(metrics['recall_at_k']) / len(metrics['recall_at_k']) if metrics['recall_at_k'] else 0,
            'mean_average_precision': sum(metrics['average_precision']) / len(metrics['average_precision']) if metrics['average_precision'] else 0
        }
        return final_metrics

    def _calculate_average_precision(self, retrieved_ids: List[str], relevant_ids: set) -> float:
        if not relevant_ids:
            return 0.0
        precisions = []
        relevant_count = 0
        for i, doc_id in enumerate(retrieved_ids):
            if doc_id in relevant_ids:
                relevant_count += 1
                precision = relevant_count / (i + 1)
                precisions.append(precision)
        return sum(precisions) / len(relevant_ids) if precisions else 0.0


# код - rag_lmclient.py

import re
import aiohttp
import asyncio
from pathlib import Path
from typing import Optional, Any, Dict, List, Union
import logging

from rag_langchain_tools import enrich_context_with_tools
from rag_prompt_utils import get_prompt_parts

class LMClientException(Exception):
    """Базовое исключение LLM клиента."""
    pass

class LMClient:
    """
    Асинхронный клиент для генерации текстов через LLM API.
    Гарантирует асинхронность, контроль длины, SRP, устойчивость к ошибкам.
    """

    def __init__(
        self,
        retriever: Any,
        data_dir: Union[str, Path],
        inform_dir: Union[str, Path],
        logger: logging.Logger,
        model_url: str,
        model_name: str,
        max_tokens: int = 1024,
        max_chars: int = 2600,
        max_attempts: int = 3,
        temperature: float = 0.7,
        timeout: int = 40,
        history_lim: int = 3,
        system_msg: Optional[str] = None
    ):
        self.retriever = retriever
        self.data_dir = Path(data_dir)
        self.inform_dir = Path(inform_dir)
        self.logger = logger

        self.model_url = model_url
        self.model_name = model_name
        self.max_tokens = max_tokens
        self.max_chars = max_chars
        self.max_attempts = max_attempts
        self.temperature = temperature
        self.timeout = timeout
        self.history_lim = history_lim
        self.system_msg = system_msg or "Вы — эксперт по бровям и ресницам."

    async def generate(
        self, 
        topic: str, 
        uploadfile: Optional[str] = None
    ) -> str:
        """
        Генерирует текст по теме с использованием контекста и инструментов.
        :param topic: Тема для генерации
        :param uploadfile: путь к прикреплённому файлу (если нужен в prompt)
        :return: Ответ LLM (или строка с ошибкой)
        """
        try:
            context = await self._get_full_context(topic)
            prompt = self._build_prompt(topic, context, uploadfile)
            return await self._request_llm_with_retries(topic, prompt)
        except Exception as e:
            self.logger.error(f"Critical error in generate: {e}")
            return "[Критическая ошибка генерации]"

    async def _get_full_context(self, topic: str) -> str:
        """
        Получает и обогащает контекст по теме.
        """
        try:
            ctx = self.retriever.retrieve(topic)
            ctx = enrich_context_with_tools(topic, ctx, self.inform_dir)
            return ctx
        except Exception as e:
            self.logger.warning(f"Ошибка получения/обогащения контекста: {e}")
            return ""

    def _build_prompt(
        self, 
        topic: str, 
        context: str, 
        uploadfile: Optional[str] = None
    ) -> str:
        """
        Генерирует промпт для LLM.
        """
        try:
            return get_prompt_parts(self.data_dir, topic, context, uploadfile=uploadfile)
        except Exception as e:
            self.logger.error(f"Ошибка генерации промпта: {e}")
            prompt_file = self.data_dir / 'prompt.txt'
            if prompt_file.exists():
                prompt_template = prompt_file.read_text(encoding='utf-8')
                return prompt_template.replace('{TOPIC}', topic).replace('{CONTEXT}', context)
            else:
                return f"{topic}\n\n{context}"

    async def _request_llm_with_retries(self, topic: str, prompt: str) -> str:
        """
        Выполняет несколько попыток генерации текста через LLM API с контролем длины и истории.
        """
        messages = [
            {"role": "system", "content": self.system_msg},
            {"role": "user", "content": prompt}
        ]
        for attempt in range(self.max_attempts):
            try:
                text = await self._request_llm(messages)
                text = self._postprocess(text)
                if len(text) <= self.max_chars:
                    self.logger.info(f"Generated text length: {len(text)} chars")
                    return text
                # Если слишком длинно — добавить в историю и просить сжать
                if attempt < self.max_attempts - 1:
                    messages = self._update_history(messages, text)
                else:
                    self.logger.warning(f"Force truncating text from {len(text)} to {self.max_chars} chars")
                    return text[:self.max_chars-10] + "..."
            except asyncio.TimeoutError as e:
                self.logger.error(f"Timeout in attempt {attempt+1}: {e}")
            except Exception as e:
                self.logger.error(f"LLM request error in attempt {attempt+1}: {e}")
            await asyncio.sleep(2)
        return "[Ошибка: превышено количество попыток генерации]"

    async def _request_llm(self, messages: List[Dict[str, str]]) -> str:
        """
        Асинхронно отправляет запрос к LLM API и получает результат.
        """
        payload = {
            "model": self.model_name,
            "messages": messages,
            "temperature": self.temperature,
            "max_tokens": self.max_tokens
        }
        timeout = aiohttp.ClientTimeout(total=self.timeout)
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.post(self.model_url, json=payload) as resp:
                if resp.status != 200:
                    raise LMClientException(f"LLM API error: HTTP {resp.status}")
                data = await resp.json()
                if 'choices' not in data or not data['choices']:
                    raise LMClientException("Invalid LLM response format")
                return data['choices'][0]['message']['content'].strip()

    def _postprocess(self, text: str) -> str:
        """
        Удаляет markdown-заголовки, разделители, мусор LLM из ответа.
        """
        rules = [
            (r"(?m)^#{2,}.*$", ""),        # markdown-заголовки
            (r"(?m)^---+", ""),            # разделители
            (r"\[\[.*?\]\]\(.*?\)", ""),   # markdown-ссылки
            (r"\n{2,}", "\n"),             # множественные переводы строк
            (r"(as an ai language model|i am an ai language model|я искусственный интеллект|как искусственный интеллект)[\.,]?\s*", "", re.IGNORECASE)
        ]
        for rule in rules:
            if len(rule) == 2:
                text = re.sub(rule[0], rule[1], text)
            else:
                text = re.sub(rule[0], rule[1], text, flags=rule[2])
        return text.strip()

    def _update_history(self, messages: List[Dict[str, str]], text: str) -> List[Dict[str, str]]:
        """
        Обновляет историю сообщений для запроса к LLM (ограничивает по self.history_lim).
        """
        # Добавляем assistant/user
        history = messages[:]
        history.append({"role": "assistant", "content": text})
        history.append({
            "role": "user",
            "content": f"Текст слишком длинный ({len(text)}>{self.max_chars}), сократи до {self.max_chars} символов."
        })
        sysm, rest = history[0], history[1:]
        # Берём последние пары user/assistant (history_lim*2), не нарушая структуру
        last_msgs = []
        for m in reversed(rest):
            if len(last_msgs) >= self.history_lim * 2:
                break
            last_msgs.insert(0, m)
        return [sysm] + last_msgs
    
# код - rag_langchain_tools.py

import logging
from rag_utils import web_search, safe_eval, analyze_table
from pathlib import Path
from typing import Optional, Dict, Any, List

logger = logging.getLogger("rag_langchain_tools")
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s [%(levelname)s] [rag_langchain_tools] %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)

TOOL_KEYWORDS = {
    "web": ["найди", "поиск", "интернет", "lookup", "search", "google", "bing", "duckduckgo"],
    "calc": ["выгод", "посчит", "calculate", "profit", "выбери", "сколько", "рассчитай"],
    "table": ["таблиц", "excel", "csv", "xlsx", "анализируй", "данные", "отчет", "таблица"]
}

def tool_internet_search(query: str, num_results: int = 8) -> str:
    """
    Выполняет интернет-поиск по запросу.
    """
    logger.info(f"Вызов интернет-поиска по запросу: {query}")
    results = web_search(query, num_results=num_results)
    if not results:
        logger.warning("Интернет-поиск не дал результатов")
        return "[Интернет-поиск не дал результатов]"
    return "\n".join(results)

def tool_calculator(expr: str, variables: Optional[Dict[str, Any]] = None) -> str:
    """
    Выполняет безопасный расчет выражения.
    """
    logger.info(f"Вызов калькулятора с выражением: {expr}")
    try:
        return str(safe_eval(expr, variables=variables))
    except Exception as e:
        logger.error(f"Ошибка калькуляции: {e}")
        return f"[Ошибка калькуляции]: {e}"

def tool_table_analysis(
    table_filename: str,
    info_query: Optional[dict]=None,
    inform_dir: Optional[str]=None,
    max_rows: int = 18,
    max_cols: int = 10
) -> str:
    """
    Анализирует таблицу по запросу.
    """
    logger.info(f"Анализ таблицы: {table_filename}")
    try:
        file_path = Path(inform_dir) / table_filename
        return analyze_table(file_path, info_query, max_rows=max_rows, max_cols=max_cols)
    except Exception as e:
        logger.error(f"Ошибка анализа таблицы: {e}")
        return f"[Ошибка анализа таблицы]: {e}"

def smart_tool_selector(
    topic: str,
    context: str,
    inform_dir: str,
    tool_keywords: Optional[Dict[str, List[str]]] = None,
    tool_log: Optional[List[str]] = None,
    max_tool_results: int = 8
) -> str:
    """
    Интеллектуальный селектор инструментов: анализирует промт, вызывает нужные инструменты,
    поддерживает цепочки, логирует действия, расширяем по ключевым словам.
    """
    tool_keywords = tool_keywords or TOOL_KEYWORDS
    tool_log = tool_log or []
    topic_lc = topic.lower()
    results = []
    used_tools = []

    # Сначала web search
    if any(x in topic_lc for x in tool_keywords["web"]):
        logger.info("[smart_tool_selector] Web search triggered")
        tool_log.append("web_search")
        results.append("[Интернет]:\n" + tool_internet_search(topic, num_results=max_tool_results))
        used_tools.append("web_search")
    # Затем калькулятор
    if any(x in topic_lc for x in tool_keywords["calc"]):
        import re
        logger.info("[smart_tool_selector] Calculator triggered")
        tool_log.append("calculator")
        m = re.search(r"(посчитай|calculate|выгоднее|выгодность|сколько)[^\d]*(.+)", topic_lc)
        expr = m.group(2) if m else topic
        results.append("[Калькулятор]:\n" + tool_calculator(expr))
        used_tools.append("calculator")
    # Затем таблица
    if any(x in topic_lc for x in tool_keywords["table"]):
        logger.info("[smart_tool_selector] Table analysis triggered")
        tool_log.append("analyze_table")
        table_files = [f.name for f in Path(inform_dir).glob("*.csv")] + [f.name for f in Path(inform_dir).glob("*.xlsx")]
        if table_files:
            results.append("[Таблица]:\n" + tool_table_analysis(table_files[0], None, inform_dir))
            used_tools.append("analyze_table")
        else:
            results.append("[Нет подходящих таблиц для анализа]")

    if used_tools:
        logger.info(f"Вызваны инструменты: {used_tools}")
    if results:
        return "\n\n".join(results)
    else:
        logger.info("Ни один инструмент не был вызван")
        return ""

def enrich_context_with_tools(
    topic: str,
    context: str,
    inform_dir: str,
    max_tool_results: int = 8
) -> str:
    """
    Добавляет к контексту максимум информации из инструментов для нейросети.
    """
    logger.info("Расширение контекста инструментами...")
    tool_result = smart_tool_selector(topic, context, inform_dir, max_tool_results=max_tool_results)
    if tool_result:
        context = context + "\n\n[Инструментальное расширение]:\n" + tool_result
        logger.info("Контекст расширен инструментами.")
    else:
        logger.info("Инструменты не были использованы для расширения контекста.")
    return context

# код - utils/config_manager.py
from pathlib import Path
import json
from typing import Any, Dict
import logging

class ConfigManager:
    """
    Менеджер конфигурации для всей системы.
    Гарантирует загрузку, валидацию и выдачу параметров с проверкой на ошибки.
    """

    def __init__(self, config_path: Path):
        """
        :param config_path: Путь к json-файлу конфигурации
        """
        self.config_path = config_path
        self.logger = logging.getLogger("config_manager")
        self.config: Dict[str, Any] = self._load_config()
        self._validate_config()

    def _load_config(self) -> Dict[str, Any]:
        """Загрузка конфигурации из файла"""
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            self.logger.critical(f"Failed to load configuration: {e}")
            raise

    def _validate_config(self) -> None:
        """Валидация конфигурации"""
        required_sections = ['telegram', 'language_model', 'retrieval', 'system', 'paths']
        for section in required_sections:
            if section not in self.config:
                raise ValueError(f"Missing required config section: {section}")

    def get_telegram_config(self) -> Dict[str, Any]:
        """
        Получение и проверка конфигурации Telegram.
        :return: словарь с ключами 'bot_token', 'channel_id' и настройками Telegram
        :raises: ValueError, если файлы токена или channel_id не найдены
        """
        config = self.config['telegram'].copy()
        try:
            token_file = Path(config['bot_token_file'])
            channel_file = Path(config['channel_id_file'])

            if not token_file.exists():
                raise ValueError(f"Telegram token file not found: {token_file}")
            if not channel_file.exists():
                raise ValueError(f"Channel ID file not found: {channel_file}")

            config['bot_token'] = token_file.read_text(encoding='utf-8').strip()
            config['channel_id'] = channel_file.read_text(encoding='utf-8').strip()
        except Exception as e:
            self.logger.critical(f"Failed to load Telegram credentials: {e}")
            raise
        return config

    def get_llm_config(self) -> Dict[str, Any]:
        """Получение настроек LLM"""
        return self.config['language_model']

    def get_retrieval_config(self) -> Dict[str, Any]:
        """Получение retrieval-конфига"""
        return self.config['retrieval']

    def get_system_config(self) -> Dict[str, Any]:
        """Получение системных параметров"""
        return self.config['system']

    def get_path(self, path_key: str) -> Path:
        """
        Получение пути из конфигурации.
        :param path_key: Ключ из секции 'paths'
        :return: Path-объект
        :raises KeyError: если ключ не найден
        """
        if path_key not in self.config['paths']:
            raise KeyError(f"Path not found in config: {path_key}")
        return Path(self.config['paths'][path_key])

    def get(self, section: str, key: str, default: Any = None) -> Any:
        """
        Получение значения по секции и ключу.
        :param section: Имя секции (str)
        :param key: Имя ключа (str)
        :param default: Значение по умолчанию
        """
        return self.config.get(section, {}).get(key, default)

    def update(self, section: str, key: str, value: Any) -> None:
        """
        Обновление значения в конфиге с сохранением в файл.
        """
        if section not in self.config:
            self.config[section] = {}
        self.config[section][key] = value
        self._save_config()

    def _save_config(self) -> None:
        """Сохранение конфигурации в файл"""
        try:
            with open(self.config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
        except Exception as e:
            self.logger.error(f"Failed to save configuration: {e}")
            raise

# код -  utils/exceptions.py
class RAGException(Exception):
    """Базовый класс для исключений RAG системы"""
    pass

class ConfigurationError(RAGException):
    """Ошибки конфигурации"""
    pass

class InitializationError(RAGException):
    """Ошибки инициализации компонентов"""
    pass

class ProcessingError(RAGException):
    """Ошибки обработки данных"""
    pass

class ModelError(RAGException):
    """Ошибки языковой модели"""
    pass

class TelegramError(RAGException):
    """Ошибки взаимодействия с Telegram"""
    pass

class FileOperationError(RAGException):
    """Ошибки файловых операций"""
    pass


# код - utils/state_manager.py
from pathlib import Path
import json
from typing import Dict, Any, Set, Optional
from datetime import datetime
import logging

class StateManager:
    """
    Менеджер состояния обработки тем.
    Сохраняет прогресс (обработанные и неудачные темы), статистику, поддерживает recoverability.
    """

    def __init__(self, state_file: Path):
        """
        :param state_file: Путь к JSON-файлу состояния.
        """
        self.state_file = state_file
        self.logger = logging.getLogger("state_manager")
        self.state: Dict[str, Any] = self._load_state()

    def _load_state(self) -> Dict[str, Any]:
        """Загрузка состояния из файла или создание по умолчанию."""
        if self.state_file.exists():
            try:
                with open(self.state_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Преобразуем processed_topics к set для быстрого поиска
                    if isinstance(data.get("processed_topics"), list):
                        data["processed_topics"] = set(data["processed_topics"])
                    elif isinstance(data.get("processed_topics"), set):
                        pass
                    else:
                        data["processed_topics"] = set()
                    return data
            except Exception as e:
                self.logger.error(f"Failed to load state: {e}")
                return self._create_default_state()
        return self._create_default_state()

    def _create_default_state(self) -> Dict[str, Any]:
        """Создание состояния по умолчанию."""
        return {
            "last_update": datetime.utcnow().isoformat(),
            "processed_topics": set(),
            "failed_topics": {},
            "statistics": {
                "total_processed": 0,
                "successful": 0,
                "failed": 0
            }
        }

    def save_state(self) -> None:
        """Сохранение состояния (atomic)."""
        try:
            state_copy = self.state.copy()
            # Сериализуем set как list для JSON
            state_copy["processed_topics"] = list(self.state["processed_topics"])
            state_copy["last_update"] = datetime.utcnow().isoformat()
            tmp_file = self.state_file.with_suffix('.tmp')
            with open(tmp_file, 'w', encoding='utf-8') as f:
                json.dump(state_copy, f, indent=4, ensure_ascii=False)
            tmp_file.replace(self.state_file)
        except Exception as e:
            self.logger.error(f"Failed to save state: {e}")
            raise

    def add_processed_topic(self, topic: str) -> None:
        """Добавить обработанную тему."""
        self.state["processed_topics"].add(topic)
        self.state["statistics"]["total_processed"] += 1
        self.state["statistics"]["successful"] += 1
        self.save_state()

    def add_failed_topic(self, topic: str, error: str) -> None:
        """Добавить тему с ошибкой."""
        self.state["failed_topics"][topic] = {
            "error": error,
            "timestamp": datetime.utcnow().isoformat(),
            "attempts": self.state["failed_topics"].get(topic, {}).get("attempts", 0) + 1
        }
        self.state["statistics"]["failed"] += 1
        self.save_state()

    def get_processed_topics(self) -> Set[str]:
        """Получить множество обработанных тем."""
        return self.state["processed_topics"]

    def get_failed_topics(self) -> Dict[str, Dict[str, Any]]:
        """Получить dict тем с ошибками."""
        return self.state["failed_topics"]

    def get_statistics(self) -> Dict[str, int]:
        """Получить статистику обработки."""
        return self.state["statistics"]

    def clear_failed_topics(self) -> None:
        """Очистить список тем с ошибками."""
        self.state["failed_topics"].clear()
        self.save_state()

# код - utils/file_manager.py

import os
import json
import csv
import requests
from typing import List, Dict, Optional, Any, Callable
from pathlib import Path
import asyncio
import aiohttp
from dataclasses import dataclass
from utils.path_utils import validate_path
from utils.exceptions import ProcessingError, FileOperationError
import time
import hashlib

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from advanced_rag_pipeline import Document, AdvancedRAGPipeline

MAX_DOCS = 10000
MAX_FILE_SIZE_MB = 50
MAX_PDF_PAGES = 200
MAX_DOCX_PARAGRAPHS = 10000

def safe_ingest(fn: Callable) -> Callable:
    """Декоратор для безопасного ingestion с логированием и единым форматом ошибок."""
    def wrapper(self, *args, **kwargs):
        try:
            return fn(self, *args, **kwargs)
        except Exception as e:
            raise FileOperationError(f"[{fn.__name__}] {e}") from e
    return wrapper

class DataIngestionManager:
    """Менеджер для загрузки и валидации данных из различных источников для RAG."""

    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.allowed_data_dir = Path("./data").resolve()

    @safe_ingest
    def load_from_text_file(self, filepath: str, encoding: str = 'utf-8') -> str:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".txt"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        with open(path, 'r', encoding=encoding) as f:
            text = f.read()
            if not text.strip():
                raise FileOperationError(f"Файл пустой или содержит только пробелы: {filepath}")
            return text

    @safe_ingest
    def load_from_pdf(self, filepath: str, max_pages: int = MAX_PDF_PAGES) -> str:
        if not PDF_AVAILABLE:
            raise ProcessingError("PyPDF2 не установлен. Установите: pip install PyPDF2")
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".pdf"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        text = ""
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                if i >= max_pages:
                    break
                text += (page.extract_text() or "") + "\n"
        if not text.strip():
            raise FileOperationError(f"PDF пустой или нечитабельный: {filepath}")
        return text

    @safe_ingest
    def load_from_docx(self, filepath: str, max_paragraphs: int = MAX_DOCX_PARAGRAPHS) -> str:
        if not DOCX_AVAILABLE:
            raise ProcessingError("python-docx не установлен. Установите: pip install python-docx")
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".docx"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        doc = docx.Document(path)
        paragraphs = doc.paragraphs[:max_paragraphs]
        text = "\n".join(paragraph.text for paragraph in paragraphs)
        if not text.strip():
            raise FileOperationError(f"DOCX пустой: {filepath}")
        return text

    @safe_ingest
    def load_from_csv(self, filepath: str, text_columns: List[str], max_rows: int = MAX_DOCS) -> List[Dict[str, Any]]:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".csv"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        documents = []
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                content_parts = [str(row[col]) for col in text_columns if col in row and row[col]]
                if content_parts:
                    content = " ".join(content_parts)
                    metadata = {k: v for k, v in row.items() if k not in text_columns}
                    documents.append({'id': f"csv_row_{i}", 'content': content, 'metadata': metadata})
        if not documents:
            raise FileOperationError(f"CSV не содержит нужных данных по столбцам {text_columns}: {filepath}")
        return documents

    @safe_ingest
    def load_from_json(self, filepath: str, content_field: str, id_field: Optional[str] = None, max_docs: int = MAX_DOCS) -> List[Dict[str, Any]]:
        path = Path(filepath)
        is_valid, reason = validate_path(path, allowed_dir=self.allowed_data_dir, allowed_exts={".json"}, max_size_mb=MAX_FILE_SIZE_MB)
        if not is_valid:
            raise FileOperationError(f"Невалидный путь/файл: {reason}")
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        documents = []
        if isinstance(data, list):
            for i, item in enumerate(data):
                if i >= max_docs:
                    break
                if content_field in item:
                    doc_id = item.get(id_field, f"json_item_{i}") if id_field else f"json_item_{i}"
                    content = str(item[content_field])
                    metadata = {k: v for k, v in item.items() if k not in [content_field, id_field]}
                    documents.append({'id': doc_id, 'content': content, 'metadata': metadata})
        if not documents:
            raise FileOperationError(f"JSON не содержит элементов с полем {content_field}: {filepath}")
        return documents

    @safe_ingest
    def load_from_url(self, url: str, timeout: int = 30) -> str:
        response = requests.get(url, timeout=timeout)
        response.raise_for_status()
        content = response.text
        if 'text/html' in response.headers.get('content-type', '') and BS4_AVAILABLE:
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            content = soup.get_text()
        if not content.strip():
            raise ProcessingError(f"URL {url} не содержит текста")
        return content

    async def load_from_urls_async(self, urls: List[str], timeout: int = 30) -> List[Dict[str, Any]]:
        async def fetch_url(session, url):
            try:
                async with session.get(url, timeout=timeout) as response:
                    content = await response.text()
                    if 'text/html' in response.headers.get('content-type', '') and BS4_AVAILABLE:
                        soup = BeautifulSoup(content, 'html.parser')
                        for script in soup(["script", "style"]):
                            script.decompose()
                        content = soup.get_text()
                    return {
                        'id': hashlib.md5(url.encode()).hexdigest(),
                        'content': content,
                        'metadata': {'source_url': url, 'status': 'success'}
                    }
            except Exception as e:
                return {
                    'id': hashlib.md5(url.encode()).hexdigest(),
                    'content': "",
                    'metadata': {'source_url': url, 'status': 'error', 'error': str(e)}
                }
        results = []
        async with aiohttp.ClientSession() as session:
            tasks = [fetch_url(session, url) for url in urls]
            fetch_results = await asyncio.gather(*tasks)
            # Сохраняем все результаты, включая ошибки
            return fetch_results

# --- Аналитика и Web-интерфейс (оставлен без изменений, кроме type hints и лимитов) ---

class RAGAnalytics:
    """Аналитика и мониторинг RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline
        self.query_log: List[Dict[str, Any]] = []

    def log_query(self, query: str, results_count: int, processing_time: float) -> None:
        self.query_log.append({
            'timestamp': time.time(),
            'query': query,
            'results_count': results_count,
            'processing_time': processing_time
        })

    def get_query_stats(self) -> Dict[str, Any]:
        if not self.query_log:
            return {"message": "Нет данных по запросам"}
        processing_times = [log['processing_time'] for log in self.query_log]
        results_counts = [log['results_count'] for log in self.query_log]
        return {
            'total_queries': len(self.query_log),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'max_processing_time': max(processing_times),
            'min_processing_time': min(processing_times),
            'avg_results_count': sum(results_counts) / len(results_counts),
            'queries_per_hour': len([q for q in self.query_log if time.time() - q['timestamp'] < 3600])
        }

    def analyze_collection_content(self, max_categories: int = 20, max_languages: int = 10) -> Dict[str, Any]:
        try:
            all_data = self.rag.collection.get()
            documents = all_data.get('documents', [])
            metadatas = all_data.get('metadatas', [])
            if not documents:
                return {"message": "Коллекция пуста"}
            total_docs = len(documents)
            total_chars = sum(len(doc) for doc in documents)
            avg_doc_length = total_chars / total_docs
            categories = {}
            languages = {}
            for metadata in metadatas:
                if metadata:
                    if 'category' in metadata:
                        cat = metadata['category']
                        if cat in categories: categories[cat] += 1
                        elif len(categories) < max_categories: categories[cat] = 1
                    if 'language' in metadata:
                        lang = metadata['language']
                        if lang in languages: languages[lang] += 1
                        elif len(languages) < max_languages: languages[lang] = 1
            return {
                'total_documents': total_docs,
                'total_characters': total_chars,
                'average_document_length': avg_doc_length,
                'categories_distribution': categories,
                'languages_distribution': languages,
                'longest_document': max(len(doc) for doc in documents),
                'shortest_document': min(len(doc) for doc in documents)
            }
        except Exception as e:
            return {"error": f"Ошибка анализа: {e}"}

class RAGWebInterface:
    """Простой веб-интерфейс для RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline, analytics: RAGAnalytics):
        self.rag = rag_pipeline
        self.analytics = analytics

    def generate_html_interface(self) -> str:
        # (оставлен без изменений, так как это статичный шаблон)
        html_template = """ ... (см. исходник) ... """
        return html_template

    def save_interface(self, filepath: str = "rag_interface.html") -> None:
        html_content = self.generate_html_interface()
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"HTML интерфейс сохранен в {filepath}")

class RAGBenchmarking:
    """Бенчмаркинг и тестирование RAG системы"""
    def __init__(self, rag_pipeline: AdvancedRAGPipeline):
        self.rag = rag_pipeline

    def create_test_dataset(self, size: int = 100) -> List[Document]:
        test_topics = [
            "машинное обучение", "искусственный интеллект", "нейронные сети",
            "векторные базы данных", "обработка естественного языка",
            "компьютерное зрение", "глубокое обучение", "анализ данных",
            "программирование на Python", "веб-разработка"
        ]
        documents = []
        for i in range(size):
            topic = test_topics[i % len(test_topics)]
            content = self._generate_synthetic_content(topic, i)
            doc = Document(
                id=f"test_doc_{i}",
                content=content,
                metadata={
                    "topic": topic,
                    "test_id": i,
                    "synthetic": True
                }
            )
            documents.append(doc)
        return documents

    def _generate_synthetic_content(self, topic: str, doc_id: int) -> str:
        templates = [
            f"{topic} является важной областью современной технологии. Документ номер {doc_id} содержит подробную информацию о применении {topic} в различных сферах.",
            f"В этом документе рассматриваются основные принципы {topic}. Это руководство номер {doc_id} поможет понять ключевые концепции.",
            f"Практическое применение {topic} демонстрируется в примере {doc_id}. Здесь представлены лучшие практики и методы."
        ]
        base_content = templates[doc_id % len(templates)]
        additional = f" Дополнительные детали включают технические аспекты, примеры использования и рекомендации экспертов в области {topic}."
        return base_content + additional

    def benchmark_search_performance(self, queries: List[str], n_runs: int = 10) -> Dict[str, Any]:
        results = {
            'queries': [],
            'avg_time': 0,
            'total_time': 0,
            'fastest_query': None,
            'slowest_query': None
        }
        total_time = 0
        fastest_time = float('inf')
        slowest_time = 0
        for query in queries:
            query_times = []
            for _ in range(n_runs):
                start_time = time.time()
                self.rag.search(query, n_results=5)
                end_time = time.time()
                query_time = end_time - start_time
                query_times.append(query_time)
                total_time += query_time
            avg_query_time = sum(query_times) / len(query_times)
            min_query_time = min(query_times)
            max_query_time = max(query_times)
            if min_query_time < fastest_time:
                fastest_time = min_query_time
                results['fastest_query'] = query
            if max_query_time > slowest_time:
                slowest_time = max_query_time
                results['slowest_query'] = query
            results['queries'].append({
                'query': query,
                'avg_time': avg_query_time,
                'min_time': min_query_time,
                'max_time': max_query_time,
                'runs': n_runs
            })
        results['avg_time'] = total_time / (len(queries) * n_runs)
        results['total_time'] = total_time
        results['fastest_time'] = fastest_time
        results['slowest_time'] = slowest_time
        return results

    def evaluate_retrieval_quality(self, test_queries: List[Dict[str, Any]]) -> Dict[str, Any]:
        metrics = {
            'precision_at_k': [],
            'recall_at_k': [],
            'average_precision': []
        }
        for test_case in test_queries:
            query = test_case['query']
            relevant_ids = set(test_case['relevant_doc_ids'])
            results = self.rag.search(query, n_results=10)
            retrieved_ids = [r.document_id for r in results]
            for k in [1, 3, 5, 10]:
                if k <= len(retrieved_ids):
                    retrieved_k = set(retrieved_ids[:k])
                    precision = len(retrieved_k & relevant_ids) / k
                    recall = len(retrieved_k & relevant_ids) / len(relevant_ids) if relevant_ids else 0
                    metrics['precision_at_k'].append(precision)
                    metrics['recall_at_k'].append(recall)
            ap = self._calculate_average_precision(retrieved_ids, relevant_ids)
            metrics['average_precision'].append(ap)
        final_metrics = {
            'mean_precision_at_k': sum(metrics['precision_at_k']) / len(metrics['precision_at_k']) if metrics['precision_at_k'] else 0,
            'mean_recall_at_k': sum(metrics['recall_at_k']) / len(metrics['recall_at_k']) if metrics['recall_at_k'] else 0,
            'mean_average_precision': sum(metrics['average_precision']) / len(metrics['average_precision']) if metrics['average_precision'] else 0
        }
        return final_metrics

    def _calculate_average_precision(self, retrieved_ids: List[str], relevant_ids: set) -> float:
        if not relevant_ids:
            return 0.0
        precisions = []
        relevant_count = 0
        for i, doc_id in enumerate(retrieved_ids):
            if doc_id in relevant_ids:
                relevant_count += 1
                precision = relevant_count / (i + 1)
                precisions.append(precision)
        return sum(precisions) / len(relevant_ids) if precisions else 0.0

# код - path_utils.py

from pathlib import Path
from typing import Optional, Set

def validate_path(
    path: Path,
    allowed_dir: Path,
    allowed_exts: Optional[Set[str]] = None,
    max_size_mb: int = 100,
    check_symlink: bool = True
) -> (bool, str):
    try:
        path = path.resolve(strict=True)
        allowed_dir = allowed_dir.resolve(strict=True)
        if check_symlink and path.is_symlink():
            return False, "Файл является symlink"
        if not (allowed_dir in path.parents or path == allowed_dir):
            return False, "Файл вне разрешённой директории"
        if allowed_exts and path.suffix.lower() not in allowed_exts:
            return False, f"Недопустимое расширение: {path.suffix}"
        if path.stat().st_size > max_size_mb * 1024 * 1024:
            return False, f"Файл слишком большой (> {max_size_mb} МБ)"
        return True, "OK"
    except Exception as e:
        return False, f"Ошибка валидации пути: {e}"
